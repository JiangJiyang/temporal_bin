from __future__ import annotations

import argparse
import json
import logging
import re
import shutil
import subprocess
import sys
import tempfile
import urllib.error
import urllib.parse
import urllib.request
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path, PurePosixPath
from typing import Any

try:
    import docx
except Exception:
    docx = None

try:
    import docx2txt
except Exception:
    docx2txt = None

try:
    import fitz
except Exception:
    fitz = None

try:
    import numpy as np
except Exception:
    np = None

try:
    import pythoncom
except Exception:
    pythoncom = None

try:
    import win32com.client
except Exception:
    win32com = None

try:
    import xlrd
except Exception:
    xlrd = None

try:
    from minio import Minio
except Exception:
    Minio = None

try:
    from openpyxl import load_workbook
except Exception:
    load_workbook = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from rapidocr_onnxruntime import RapidOCR
except Exception:
    RapidOCR = None


PROJECT_ROOT = Path(__file__).resolve().parent
INPUT_FIELDS = ("zipName", "folderName", "fileName", "fileFetchPath")
OUTPUT_FIELDS = ("unit", "product", "project_type", "project_name", "project_stage")
PROJECT_TYPE_VALUES = ("固定资产投资", "科研项目", "规划", "政策法规", "其他")
PROJECT_STAGE_VALUES = ("立项", "可研", "初设", "竣工")
GENERIC_SEGMENTS = {
    "测试项目",
    "文件",
    "附件",
    "专家意见回复",
    "建筑",
    "结构",
    "暖通",
    "工艺",
    "电气",
    "机电公司",
    "一所",
    "总图",
    "给排水",
    "水",
    "动力",
}
ORG_MARKERS = (
    "有限公司",
    "有限责任公司",
    "股份有限公司",
    "集团公司",
    "集团",
    "研究院",
    "研究所",
    "设计院",
    "设计研究院",
    "中心",
    "事务所",
    "委员会",
    "实验室",
    "学校",
    "学院",
    "大学",
    "公司",
    "厂",
    "所",
)
ORG_LEGAL_SUFFIXES = (
    "有限责任公司",
    "股份有限公司",
    "有限公司",
    "集团公司",
    "设计研究院",
    "研究院",
    "研究所",
    "设计院",
    "中心",
    "学校",
    "学院",
    "大学",
    "公司",
    "厂",
    "所",
    "集团",
)
UNIT_LABEL_WEIGHTS = {
    "建设单位": 16,
    "项目单位": 15,
    "项目法人": 15,
    "业主单位": 15,
    "建设管理单位": 15,
    "申报单位": 14,
    "承担单位": 13,
    "编制单位": 10,
    "设计单位": 9,
    "编制机构": 9,
}
STAGE_KEYWORDS = {
    "立项": ("项目建议书", "建议书", "立项", "备案", "申请报告", "项目申请", "申报", "批复"),
    "可研": ("可行性研究报告", "可行性研究", "可研", "研究报告", "评估报告", "实施方案", "方案"),
    "初设": ("初步设计", "初设", "施工图", "设计说明", "设计概算"),
    "竣工": ("竣工验收报告", "竣工验收总结报告", "竣工验收", "竣工总结", "竣工", "完工", "后评价", "验收"),
}
PROJECT_TYPE_KEYWORDS = {
    "政策法规": ("办法", "规定", "条例", "通知", "规范", "制度", "细则", "通告", "公告"),
    "规划": ("规划", "总体规划", "专项规划", "布局规划", "发展规划", "建设规划"),
    "科研项目": ("科研", "课题", "研究课题", "实验", "试验", "研发", "科研项目"),
    "固定资产投资": ("固定资产投资", "投资项目", "建设项目", "工程项目", "技术改造项目", "技改项目", "改造项目", "厂房", "车间", "可行性研究", "初步设计", "竣工验收", "项目建议书", "概算"),
}
PROJECT_SUFFIX_PATTERNS = (
    r"项目建议书$",
    r"建议书$",
    r"可行性研究报告$",
    r"可行性研究$",
    r"初步设计$",
    r"初设$",
    r"竣工验收总结报告$",
    r"竣工验收报告$",
    r"竣工验收$",
    r"专项规划$",
    r"总体规划$",
    r"规划$",
    r"编制说明$",
    r"评估报告$",
    r"设备论证报告$",
    r"意见回复$",
    r"回复$",
    r"报告$",
    r"正文$",
    r"目录$",
    r"封面$",
    r"标签$",
    r"归档$",
    r"发文$",
)
TITLE_HINT_RE = re.compile(r"([^\n]{4,120}?(?:项目|工程|规划|办法|条例))(?:建议书|项目建议书|可行性研究报告|初步设计|竣工验收总结报告|竣工验收报告|竣工验收|专项规划|规划|编制说明)?")
SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".xls", ".xlsx"}
ESCAPED_UNICODE_RE = re.compile(r"#U([0-9a-fA-F]{4})")
CONTEXT_DOC_PRIORITY = ("标签", "封面", "项目建议书", "建议书", "可行性研究", "可研", "初步设计", "初设", "竣工", "正文", "目录", "批复", "报告")
PROJECT_NAME_HINTS = ("项目", "工程", "规划", "办法", "条例", "改造", "建设", "车间", "基地", "桥梁", "大桥", "办公楼", "主管道", "核电", "设备", "船", "厂房")
PROJECT_NAME_REJECT_HINTS = (
    "项目负责人",
    "项目单位法人",
    "单位法人",
    "单位：",
    "总图编号",
    "工程量",
    "详见下表",
    "设计依据",
    "编制依据",
    "专家意见",
    "反馈表",
    "评审意见",
    "对比表",
    "项目或费用名称",
    "子项名称",
    "设计证书编号",
    "目录",
    "CLIENT",
    "发文",
    "上级批文",
)
ORG_REJECT_HINTS = ("项目名称", "项目负责人", "单位：", "详见下表", "万元", "设备", "分析", "以下简称", "目录", "设计依据", "你单位", "我单位", "贵单位", "关于", "批复", "请示", "调减", "同意", "《", "》")
ORG_FRAGMENT_RE = re.compile(r"([^\n，。；;|:：]{2,60}?(?:有限责任公司|股份有限公司|有限公司|集团公司|设计研究院|研究院|研究所|设计院|中心|学校|学院|大学|公司|厂|所|集团(?!有限公司|有限责任公司|股份有限公司)))")
CONTEXT_NEUTRAL_HINTS = ("标签", "封面", "目录")
ATTACHMENT_HINTS = ("附件", "附表", "附图")
REVIEW_HINTS = ("专家意见", "意见回复", "反馈表", "评审意见", "评估报告", "对比表", "概算", "估算", "经济分析", "蓝图汇总")
SUPPORT_DOC_HINTS = ATTACHMENT_HINTS + REVIEW_HINTS + ("设备论证报告", "专题论证报告")
UNIT_NEGATIVE_CONTEXT_HINTS = ("制造商", "供应商", "供货商", "生产厂家", "设备厂家", "联系人", "联系电话", "传真", "邮箱", "网址", "设计证书", "投标", "代理商")
ORG_SENTENCE_NOISE_HINTS = ("你单位", "我单位", "贵单位", "关于", "批复", "请示", "调减", "同意", "《", "》")
GENERIC_CONTEXT_SEGMENTS = {
    "测试项目",
    "文件",
    "附件",
    "附表",
    "附图",
    "工艺",
    "建筑",
    "结构",
    "暖通",
    "水",
    "电气",
    "动力",
    "总体",
    "技经",
    "专家意见",
    "专家意见回复",
    "意见回复",
    "机电公司",
    "一所",
    "二所",
    "三所",
    "四所",
    "五所",
}
ORG_NOISE_KEYWORDS = (
    "优先采用",
    "采用",
    "不属于",
    "属于",
    "说明",
    "补充说明",
    "文字要求",
    "文本中的位置",
    "要求",
    "增加",
    "由厂",
    "来自厂",
    "排至",
    "输送",
    "进厂",
)
ORG_NOISE_PREFIX_RE = re.compile(
    r"^(?:优先(?:采用)?|采用|拟采用|原址土地(?:已)?不属于|不属于|属于|说明|补充说明|文字要求|文本中的位置与目前|文本中的位置|文本中|要求|增加|其中|由于|由|来自|排至|输送|以前进厂|进厂时间|进厂|本次设计内容为|本工程|本次技改|本次设计|室内生活、生产给水由|生活污水排至|压缩空气管道由|天然气由)+"
)
ORG_FACILITY_SUFFIXES = ("变电所", "配电室", "主厂", "生产厂", "厂区", "车间", "厂房", "办公室", "休息间", "工具间", "建筑物")
PROJECT_NAME_STRONG_HINTS = ("项目", "工程", "规划", "办法", "条例")
PROJECT_NAME_PATH_NOISE_HINTS = (
    "归档",
    "审后修改",
    "修改",
    "反馈表",
    "评审意见",
    "对比表",
    "估算",
    "分析表",
    "设备表",
    "正文",
    "目录",
    "封面",
    "标签",
    "蓝图汇总",
    "ddd",
)
PROJECT_NAME_HEADING_HINTS = ("建设规模及项目", "主要技术经济指标和工程", "建设性质及设计范围", "建设条件", "建设目标", "项目组成")
STAGE_REFERENCE_HINTS = SUPPORT_DOC_HINTS + ("批复", "回复", "反馈")
PROJECT_TRAILING_NOISE_PATTERNS = (
    r"(?:项目建议书|建议书|可行性研究报告|可行性研究|可研|初步设计|初设|竣工验收总结报告|竣工验收报告|竣工验收|竣工|专项规划|总体规划|规划)(?:归档|上报|发文|评审|评估|审后修改|修改稿|送审稿|蓝图汇总)*$",
    r"(?:归档|上报|发文|评审|评估|审后修改|修改稿|送审稿|蓝图汇总|反馈表|意见回复|专家意见回复)$",
    r"(?:总稿|正文|目录|封面|标签|对比表|概算|估算|经济分析)$",
)


@dataclass
class ModelProvider:
    name: str
    config: dict[str, Any]


@dataclass
class MinioState:
    available: bool
    client: Any = None
    reason: str = ""
    endpoint: str = ""


@dataclass
class ExtractedDocument:
    logical_path: str
    resolved_path: Path
    text: str
    front_text: str
    used_ocr: bool
    source: str
    errors: list[str] = field(default_factory=list)
    minio_object: str = ""


@dataclass
class DocSignals:
    owner_units: Counter[str] = field(default_factory=Counter)
    support_units: Counter[str] = field(default_factory=Counter)
    explicit_project_names: Counter[str] = field(default_factory=Counter)
    general_project_names: Counter[str] = field(default_factory=Counter)
    stage_scores: Counter[str] = field(default_factory=Counter)
    type_scores: Counter[str] = field(default_factory=Counter)


class OfficeBridge:
    def __init__(self, logger: logging.Logger):
        self.logger = logger
        self.word = None
        self.excel = None
        self._coinitialized = False
        if pythoncom is not None:
            try:
                pythoncom.CoInitialize()
                self._coinitialized = True
            except Exception as exc:
                self.logger.warning("COM 初始化失败: %s", exc)

    def _ensure_word(self) -> Any:
        if self.word is None:
            if win32com is None:
                raise RuntimeError("pywin32 未安装，无法读取 doc")
            self.word = win32com.client.DispatchEx("Word.Application")
            self.word.Visible = False
            self.word.DisplayAlerts = 0
        return self.word

    def _ensure_excel(self) -> Any:
        if self.excel is None:
            if win32com is None:
                raise RuntimeError("pywin32 未安装，无法读取 Excel")
            self.excel = win32com.client.DispatchEx("Excel.Application")
            self.excel.Visible = False
            self.excel.DisplayAlerts = False
        return self.excel

    def _run_command(self, command: list[str]) -> str:
        completed = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="ignore", check=True)
        return completed.stdout.strip()

    def _read_doc_via_external_tools(self, path: Path) -> str:
        for tool in ("antiword", "catdoc"):
            executable = shutil.which(tool)
            if not executable:
                continue
            try:
                return self._run_command([executable, str(path)])
            except Exception as exc:
                self.logger.warning("%s 读取失败: %s", tool, exc)
        for tool in ("soffice", "libreoffice"):
            executable = shutil.which(tool)
            if not executable:
                continue
            with tempfile.TemporaryDirectory(prefix="zch_doc_convert_") as temp_dir:
                try:
                    subprocess.run([executable, "--headless", "--convert-to", "txt:Text", "--outdir", temp_dir, str(path)], capture_output=True, text=True, encoding="utf-8", errors="ignore", check=True)
                    output_path = Path(temp_dir) / f"{path.stem}.txt"
                    if output_path.exists():
                        return output_path.read_text(encoding="utf-8", errors="ignore")
                except Exception as exc:
                    self.logger.warning("%s 转 txt 失败: %s", tool, exc)
        raise RuntimeError("当前环境缺少可用的 doc 读取后端")

    def _is_useful_probe_line(self, value: str) -> bool:
        text = _normalize_output_text(value)
        if len(text) < 4 or len(text) > 120:
            return False
        visible = len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", text))
        if visible < max(4, len(text) // 3):
            return False
        return any(label in text for label in UNIT_LABEL_WEIGHTS) or is_org_candidate(text) or has_project_hint(text) or bool(normalize_project_stage(text))

    def _read_doc_via_binary_probe(self, path: Path) -> str:
        try:
            data = path.read_bytes()
        except Exception:
            return ""
        lines: list[str] = []
        seen: set[str] = set()
        for encoding, offsets in (("utf-16le", (0, 1)), ("gb18030", (0,)), ("utf-8", (0,))):
            for offset in offsets:
                try:
                    decoded = data[offset:].decode(encoding, errors="ignore")
                except Exception:
                    continue
                decoded = decoded.replace("\x00", " ").replace("\ufeff", " ")
                for raw in re.split(r"[\r\n]+", decoded):
                    cleaned = _normalize_output_text(raw)
                    if cleaned in seen or not self._is_useful_probe_line(cleaned):
                        continue
                    seen.add(cleaned)
                    lines.append(cleaned)
                    if len(lines) >= 200:
                        return "\n".join(lines)
        return "\n".join(lines)

    def read_doc_text(self, path: Path) -> str:
        if win32com is not None:
            try:
                word = self._ensure_word()
                document = word.Documents.Open(str(path), False, True)
                try:
                    return str(document.Content.Text or "")
                finally:
                    document.Close(False)
            except Exception as exc:
                self.logger.warning("Word COM 读取失败，尝试外部转换器: %s", exc)
        try:
            return self._read_doc_via_external_tools(path)
        except Exception as exc:
            probe_text = self._read_doc_via_binary_probe(path)
            if probe_text:
                self.logger.warning("doc 读取降级到二进制探测: %s: %s", path.name, exc)
                return probe_text
            raise RuntimeError(f"{exc}; 且二进制探测未提取到有效文本")

    def read_excel_text(self, path: Path, max_rows: int, max_cols: int) -> str:
        excel = self._ensure_excel()
        workbook = excel.Workbooks.Open(str(path), ReadOnly=True)
        try:
            worksheet = workbook.Worksheets(1)
            row_count = min(int(worksheet.UsedRange.Rows.Count), max_rows)
            col_count = min(int(worksheet.UsedRange.Columns.Count), max_cols)
            lines = [str(worksheet.Name)]
            for row_index in range(1, row_count + 1):
                cells: list[str] = []
                for col_index in range(1, col_count + 1):
                    value = str(worksheet.Cells(row_index, col_index).Text or "").strip()
                    if value:
                        cells.append(value)
                if cells:
                    lines.append(" | ".join(cells))
            return "\n".join(lines)
        finally:
            workbook.Close(False)

    def close(self) -> None:
        if self.word is not None:
            try:
                self.word.Quit()
            except Exception:
                pass
        if self.excel is not None:
            try:
                self.excel.Quit()
            except Exception:
                pass
        self.word = None
        self.excel = None
        if self._coinitialized and pythoncom is not None:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            self._coinitialized = False

class DocumentExtractor:
    def __init__(self, config: dict[str, Any], logger: logging.Logger):
        self.config = config
        self.logger = logger
        self.office = OfficeBridge(logger)
        self._ocr_engine = None

    def close(self) -> None:
        self.office.close()

    def _clean_text(self, text: str) -> str:
        text = text or ""
        text = text.replace("\ufeff", "")
        text = text.replace("\u3000", " ")
        text = text.replace("\x07", " ")
        text = text.replace("\x13", " ").replace("\x14", " ").replace("\x15", " ")
        text = text.replace("\x0b", "\n")
        text = text.replace("\x0c", "\n")
        text = re.sub(r"[\x00-\x08\x0e-\x1f]", " ", text)
        text = re.sub(r"HYPERLINK\s+\\l\s+\"[^\"]+\"", " ", text, flags=re.IGNORECASE)
        text = re.sub(r"PAGEREF\s+[^\s]+(?:\s+\\h)?", " ", text, flags=re.IGNORECASE)
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _read_docx_text(self, path: Path) -> str:
        chunks: list[str] = []
        if docx2txt is not None:
            try:
                parsed = docx2txt.process(str(path)) or ""
                if parsed.strip():
                    chunks.append(parsed)
            except Exception as exc:
                self.logger.warning("docx2txt 读取失败: %s: %s", path.name, exc)
        if docx is not None:
            try:
                document = docx.Document(str(path))
                for paragraph in document.paragraphs:
                    if paragraph.text and paragraph.text.strip():
                        chunks.append(paragraph.text)
                for table in document.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                        if cells:
                            chunks.append(" | ".join(cells))
                for section in document.sections:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text and paragraph.text.strip():
                            chunks.append(paragraph.text)
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text and paragraph.text.strip():
                            chunks.append(paragraph.text)
            except Exception as exc:
                self.logger.warning("python-docx 读取失败: %s: %s", path.name, exc)
        return self._clean_text("\n".join(chunks))

    def _read_xlsx_text(self, path: Path, max_rows: int, max_cols: int) -> str:
        if load_workbook is None:
            return ""
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        try:
            worksheet = workbook[workbook.sheetnames[0]]
            lines = [str(worksheet.title)]
            for row in worksheet.iter_rows(min_row=1, max_row=max_rows, max_col=max_cols, values_only=True):
                cells = [_stringify_cell(value) for value in row if _stringify_cell(value)]
                if cells:
                    lines.append(" | ".join(cells))
            return self._clean_text("\n".join(lines))
        finally:
            workbook.close()

    def _read_xls_text(self, path: Path, max_rows: int, max_cols: int) -> str:
        if xlrd is None:
            return ""
        workbook = xlrd.open_workbook(str(path), on_demand=True)
        try:
            sheet = workbook.sheet_by_index(0)
            lines = [str(sheet.name)]
            for row_index in range(min(sheet.nrows, max_rows)):
                cells: list[str] = []
                for col_index in range(min(sheet.ncols, max_cols)):
                    cell = _stringify_cell(sheet.cell_value(row_index, col_index))
                    if cell:
                        cells.append(cell)
                if cells:
                    lines.append(" | ".join(cells))
            return self._clean_text("\n".join(lines))
        finally:
            workbook.release_resources()

    def _read_pdf_text(self, path: Path, max_pages: int) -> str:
        if PdfReader is None:
            return ""
        reader = PdfReader(str(path))
        chunks: list[str] = []
        for page in reader.pages[:max_pages]:
            try:
                chunks.append(page.extract_text() or "")
            except Exception as exc:
                self.logger.warning("PDF 页面读取失败: %s: %s", path.name, exc)
        return self._clean_text("\n".join(chunks))

    def _should_run_ocr(self, text: str) -> bool:
        compact_length = len(re.sub(r"\s+", "", text or ""))
        if not self.config["ocr"].get("enabled", True):
            return False
        if self.config["ocr"].get("only_when_text_empty", True):
            return compact_length == 0
        return compact_length < int(self.config["ocr"].get("min_chars_to_skip_ocr", 80))

    def _ensure_ocr_engine(self) -> Any:
        if self._ocr_engine is None:
            if RapidOCR is None:
                raise RuntimeError("RapidOCR 未安装")
            self._ocr_engine = RapidOCR()
        return self._ocr_engine

    def _ocr_pdf_text(self, path: Path, max_pages: int, dpi: int) -> str:
        if fitz is None or RapidOCR is None or np is None:
            return ""
        engine = self._ensure_ocr_engine()
        pdf = fitz.open(str(path))
        try:
            lines: list[str] = []
            scale = max(float(dpi) / 72.0, 1.0)
            for index in range(min(len(pdf), max_pages)):
                page = pdf[index]
                pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                image = np.frombuffer(pixmap.samples, dtype=np.uint8).reshape(pixmap.height, pixmap.width, pixmap.n)
                result, _ = engine(image)
                if result:
                    for item in result:
                        value = item[1] if len(item) > 1 else ""
                        if isinstance(value, (list, tuple)):
                            value = "".join(str(part) for part in value)
                        value = str(value).strip()
                        if value:
                            lines.append(value)
            return self._clean_text("\n".join(lines))
        finally:
            pdf.close()

    def extract(self, path: Path, logical_path: str) -> ExtractedDocument:
        suffix = path.suffix.lower()
        max_pages = int(self.config["ocr"].get("max_pdf_pages", 8))
        max_rows = int(self.config["ocr"].get("max_excel_rows", 30))
        max_cols = int(self.config["ocr"].get("max_excel_cols", 12))
        max_front_chars = int(self.config["ocr"].get("max_front_chars", 4000))
        errors: list[str] = []
        used_ocr = False
        text = ""
        try:
            if suffix == ".pdf":
                text = self._read_pdf_text(path, max_pages)
                if self._should_run_ocr(text):
                    self.logger.warning("OCR 触发: %s", logical_path)
                    ocr_text = self._ocr_pdf_text(path, max_pages, int(self.config["ocr"].get("dpi", 300)))
                    if ocr_text:
                        used_ocr = True
                        text = self._clean_text(f"{text}\n{ocr_text}")
            elif suffix == ".doc":
                text = self._clean_text(self.office.read_doc_text(path))
            elif suffix == ".docx":
                text = self._read_docx_text(path)
                if not text:
                    text = self._clean_text(self.office.read_doc_text(path))
            elif suffix == ".xls":
                text = self._read_xls_text(path, max_rows, max_cols)
                if not text:
                    text = self._clean_text(self.office.read_excel_text(path, max_rows, max_cols))
            elif suffix == ".xlsx":
                text = self._read_xlsx_text(path, max_rows, max_cols)
                if not text:
                    text = self._clean_text(self.office.read_excel_text(path, max_rows, max_cols))
            else:
                errors.append(f"不支持的文件类型: {suffix}")
        except Exception as exc:
            errors.append(str(exc))
        if not text:
            errors.append("文档文本为空")
        return ExtractedDocument(logical_path=logical_path, resolved_path=path, text=text, front_text=text[:max_front_chars], used_ocr=used_ocr, source="local", errors=errors)


class RuntimeContext:
    def __init__(self, config_path: str | Path, override_test_mode: bool | None = None):
        self.config_path = Path(config_path).resolve()
        self.config = load_runtime_config(self.config_path)
        if override_test_mode is not None:
            self.config["test_mode"] = override_test_mode
        self.logger = build_logger(self.config)
        self.prompt_cache: dict[str, dict[str, str]] = {}
        self.prompt = self._load_prompt(self.config["paths"]["prompt_path"])
        self.minio = check_minio(self.config["minio"], self.logger)
        self.model = choose_model_provider(self.config, self.logger)
        self.extractor = DocumentExtractor(self.config, self.logger)
        self.document_cache: dict[str, ExtractedDocument] = {}
        self.project_context_cache: dict[str, list[tuple[dict[str, str], ExtractedDocument, int]]] = {}
        self._temp_files: list[Path] = []
        if self.config["test_mode"]:
            if self.minio.available and self.config["minio"].get("prefer_remote_in_test", False):
                self.logger.info("测试模式: MinIO 可连通，优先从 MinIO 读取")
            elif self.minio.available:
                self.logger.info("测试模式: MinIO 可连通，但默认仍以本地测试目录为主")
            else:
                self.logger.warning("测试模式 MinIO 检查失败，回退本地测试目录: %s", self.minio.reason)
        elif not self.minio.available:
            if self.config["minio"].get("allow_local_fallback", True):
                self.logger.warning("生产模式 MinIO 检查失败，允许本地降级继续: %s", self.minio.reason)
            else:
                self.logger.error("生产模式 MinIO 检查失败: %s", self.minio.reason)

    def _load_prompt(self, prompt_path: str | Path) -> dict[str, str]:
        key = str(Path(prompt_path).resolve())
        if key in self.prompt_cache:
            return self.prompt_cache[key]
        payload = {
            "system": "你是文档元数据抽取器。必须严格输出 JSON，只能包含 unit、product、project_type、project_name、project_stage。",
            "user_template": "请根据路径上下文、前部文本和规则结果，输出 JSON。\nproject_type 只能是：固定资产投资、科研项目、规划、政策法规、其他。\nproject_stage 只能是：立项、可研、初设、竣工。\nproduct 固定为 N/A。\n不确定就返回空字符串。\n\n路径上下文:\n{path_context}\n\n前部文本:\n{front_text}\n\n规则结果:\n{rule_result}\n",
        }
        path = Path(prompt_path)
        if path.exists():
            try:
                loaded = json.loads(path.read_text(encoding="utf-8"))
                if isinstance(loaded, dict):
                    payload["system"] = str(loaded.get("system", payload["system"]))
                    payload["user_template"] = str(loaded.get("user_template", payload["user_template"]))
            except Exception as exc:
                self.logger.warning("prompt.json 读取失败，使用内置模板: %s", exc)
        self.prompt_cache[key] = payload
        return payload

    def resolve_prompt(self, override_path: str = "") -> dict[str, str]:
        return self.prompt if not override_path else self._load_prompt(resolve_reference_path(override_path, self.config_path.parent, PROJECT_ROOT))

    def resolve_local_path(self, file_fetch_path: str) -> Path:
        raw = (file_fetch_path or "").strip()
        if not raw:
            return Path("")
        candidate = Path(raw)
        if candidate.is_absolute():
            return candidate
        rel = PurePosixPath(raw.replace("\\", "/"))
        return Path(self.config["paths"]["data_root"]).joinpath(*rel.parts)

    def download_from_minio(self, file_fetch_path: str) -> tuple[Path, str]:
        if not self.minio.available or self.minio.client is None:
            raise FileNotFoundError(self.minio.reason or "MinIO 不可用")
        bucket = str(self.config["minio"].get("bucket", "")).strip()
        suffix = Path(file_fetch_path).suffix
        for object_name in build_minio_object_candidates(str(self.config["minio"].get("source_prefix", "") or "").strip(), file_fetch_path):
            try:
                self.minio.client.stat_object(bucket, object_name)
                file_descriptor, temp_name = tempfile.mkstemp(suffix=suffix or ".tmp")
                try:
                    import os

                    os.close(file_descriptor)
                except Exception:
                    pass
                Path(temp_name).unlink(missing_ok=True)
                temp_path = Path(temp_name)
                self.minio.client.fget_object(bucket, object_name, str(temp_path))
                self._temp_files.append(temp_path)
                return temp_path, object_name
            except Exception:
                continue
        raise FileNotFoundError(f"MinIO 中未找到对象: {file_fetch_path}")

    def resolve_document(self, item: dict[str, str]) -> ExtractedDocument:
        cache_key = item["fileFetchPath"].replace("\\", "/")
        if cache_key in self.document_cache:
            return self.document_cache[cache_key]
        if not self.config["test_mode"] and not self.minio.available:
            extracted = ExtractedDocument(logical_path=item["fileFetchPath"], resolved_path=Path(""), text="", front_text="", used_ocr=False, source="blocked", errors=[f"生产模式 MinIO 不可连通: {self.minio.reason}"])
            self.document_cache[cache_key] = extracted
            return extracted
        source = "local"
        resolved_path = Path("")
        minio_object = ""
        try:
            if self.config["test_mode"] and self.minio.available:
                resolved_path, minio_object = self.download_from_minio(item["fileFetchPath"])
                source = "minio"
            else:
                resolved_path = self.resolve_local_path(item["fileFetchPath"])
        except Exception as exc:
            if self.config["test_mode"]:
                self.logger.warning("MinIO 读取失败，回退本地文件: %s", exc)
                resolved_path = self.resolve_local_path(item["fileFetchPath"])
            else:
                extracted = ExtractedDocument(logical_path=item["fileFetchPath"], resolved_path=Path(""), text="", front_text="", used_ocr=False, source="minio", errors=[str(exc)])
                self.document_cache[cache_key] = extracted
                return extracted
        if not resolved_path.exists():
            extracted = ExtractedDocument(logical_path=item["fileFetchPath"], resolved_path=resolved_path, text="", front_text="", used_ocr=False, source=source, errors=[f"关键文件无法读取: {resolved_path}"], minio_object=minio_object)
            self.document_cache[cache_key] = extracted
            return extracted
        extracted = self.extractor.extract(resolved_path, item["fileFetchPath"])
        extracted.source = source
        extracted.minio_object = minio_object
        if source == "minio":
            self.logger.info("MinIO 读取成功: %s", minio_object)
        self.document_cache[cache_key] = extracted
        return extracted

    def _iter_local_context_items(self, root_prefix: str, zip_name: str, exclude_paths: set[str], preferred_stage: str) -> list[tuple[dict[str, str], int]]:
        root_path = self.resolve_local_path(root_prefix)
        if not root_path.exists():
            return []
        candidates: list[tuple[dict[str, str], int]] = []
        for path in root_path.rglob("*"):
            if not path.is_file():
                continue
            relative = normalize_path_text(str(path.relative_to(Path(self.config["paths"]["data_root"]))))
            if relative in exclude_paths or not is_supported_extension(relative):
                continue
            score = score_context_path(relative, preferred_stage)
            if score <= 0:
                continue
            candidates.append((build_context_item(zip_name, relative), score))
        return candidates

    def _iter_minio_context_items(self, root_prefix: str, zip_name: str, exclude_paths: set[str], preferred_stage: str) -> list[tuple[dict[str, str], int]]:
        if not self.minio.available or self.minio.client is None:
            return []
        bucket = str(self.config["minio"].get("bucket", "")).strip()
        candidates: list[tuple[dict[str, str], int]] = []
        seen: set[str] = set()
        for prefix in build_minio_object_candidates(str(self.config["minio"].get("source_prefix", "") or "").strip(), root_prefix):
            try:
                objects = self.minio.client.list_objects(bucket, prefix=prefix.strip("/"), recursive=True)
            except Exception:
                continue
            for obj in objects:
                object_name = normalize_path_text(getattr(obj, "object_name", "")).strip("/")
                if not object_name or object_name in seen:
                    continue
                seen.add(object_name)
                if object_name in exclude_paths or not is_supported_extension(object_name):
                    continue
                score = score_context_path(object_name, preferred_stage)
                if score <= 0:
                    continue
                candidates.append((build_context_item(zip_name, object_name), score))
        return candidates

    def discover_project_context(self, payload: dict[str, list[str]]) -> list[tuple[dict[str, str], ExtractedDocument, int]]:
        root_prefix = payload_project_root(payload)
        if not root_prefix:
            return []
        preferred_stage = normalize_project_stage(" ".join(payload.get("folderName", []) + payload.get("fileName", []) + payload.get("zipName", [])))
        zip_name = payload.get("zipName", [""])[0] if payload.get("zipName") else ""
        exclude_paths = {normalize_path_text(path).strip("/") for path in payload.get("fileFetchPath", [])}
        cache_key = f"{root_prefix}|{preferred_stage}|{'|'.join(sorted(exclude_paths))}"
        if cache_key in self.project_context_cache:
            return self.project_context_cache[cache_key]
        candidates = self._iter_minio_context_items(root_prefix, zip_name, exclude_paths, preferred_stage) if self.config["test_mode"] and self.minio.available else self._iter_local_context_items(root_prefix, zip_name, exclude_paths, preferred_stage)
        candidates.sort(key=lambda item: (item[1], len(item[0]["fileFetchPath"])), reverse=True)
        selected: list[tuple[dict[str, str], ExtractedDocument, int]] = []
        role_limits = {"anchor": 2, "main": 2, "supplement": 1}
        role_counts: Counter[str] = Counter()
        seen_selected_paths: set[str] = set()
        for item, score in candidates:
            logical_path = item["fileFetchPath"]
            if logical_path in seen_selected_paths:
                continue
            role = "anchor" if is_anchor_document(logical_path) else "main" if is_main_project_doc(logical_path, preferred_stage) else "supplement"
            if role_counts[role] >= role_limits[role]:
                continue
            extracted = self.resolve_document(item)
            if extracted.errors:
                continue
            role_counts[role] += 1
            seen_selected_paths.add(logical_path)
            weight = 4 if role == "anchor" else 3 if role == "main" else 2
            selected.append((item, extracted, weight))
            if len(selected) >= sum(role_limits.values()):
                break
        if selected:
            self.logger.info("补充项目上下文: %s -> %s", root_prefix, ", ".join(item["fileName"] for item, _, _ in selected))
        self.project_context_cache[cache_key] = selected
        return selected

    def maybe_refine_with_llm(self, payload: dict[str, list[str]], batch_result: dict[str, str], documents: list[ExtractedDocument], prompt_override: str = "") -> dict[str, str]:
        if self.model is None:
            return batch_result
        needs_refine = not all(batch_result.get(field) for field in ("unit", "project_type", "project_name", "project_stage"))
        needs_refine = needs_refine or looks_like_noise_title(batch_result.get("project_name", "")) or len(batch_result.get("project_name", "")) < 6
        needs_refine = needs_refine or has_org_noise(batch_result.get("unit", "")) or bool(re.match(r"^\d", batch_result.get("unit", "")))
        path_stage = normalize_project_stage(" ".join(payload.get("folderName", []) + payload.get("fileName", []) + payload.get("zipName", [])))
        needs_refine = needs_refine or bool(path_stage and batch_result.get("project_stage") and batch_result["project_stage"] != path_stage)
        if not needs_refine:
            return batch_result
        prompt = self.resolve_prompt(prompt_override)
        top_documents = sorted(documents, key=lambda doc: (_doc_priority(doc.logical_path), len(doc.front_text)), reverse=True)[:3]
        path_context = json.dumps({field: payload[field][:3] for field in INPUT_FIELDS}, ensure_ascii=False, indent=2)
        front_text = "\n\n".join(f"[{doc.logical_path}]\n{doc.front_text[:3000]}" for doc in top_documents)
        user_prompt = render_prompt_template(
            prompt["user_template"],
            path_context=path_context,
            front_text=front_text,
            rule_result=json.dumps(batch_result, ensure_ascii=False),
        )
        try:
            content = call_openai_compatible(self.model, prompt["system"], user_prompt)
            parsed = parse_json_object(content)
            if not isinstance(parsed, dict):
                return batch_result
            merged = dict(batch_result)
            merged["unit"] = _normalize_output_text(parsed.get("unit")) or batch_result["unit"]
            merged["product"] = "N/A"
            merged["project_type"] = normalize_project_type(_normalize_output_text(parsed.get("project_type"))) or batch_result["project_type"]
            merged["project_name"] = clean_project_name(_normalize_output_text(parsed.get("project_name")), merged["unit"]) or batch_result["project_name"]
            merged["project_stage"] = normalize_project_stage(_normalize_output_text(parsed.get("project_stage"))) or batch_result["project_stage"]
            return merged
        except Exception as exc:
            self.logger.warning("模型补全失败，使用规则结果: %s", exc)
            self.model = None
            return batch_result

    def close(self) -> None:
        self.extractor.close()
        for path in self._temp_files:
            try:
                path.unlink(missing_ok=True)
            except Exception:
                pass
        self._temp_files.clear()

def _as_bool(value: Any, default: bool) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        lowered = value.strip().lower()
        if lowered in {"1", "true", "yes", "y", "on"}:
            return True
        if lowered in {"0", "false", "no", "n", "off"}:
            return False
    return default


def _as_int(value: Any, default: int) -> int:
    try:
        return int(value)
    except Exception:
        return default


def _stringify_cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _normalize_output_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def normalize_for_compare(value: str) -> str:
    return re.sub(r"[\s\-_./\\()（）【】\[\]<>《》“”\"'`·,，。:：;；]+", "", value or "").lower()


def normalize_path_text(value: str) -> str:
    return (value or "").replace("\\", "/")


def contains_any(text: str, keywords: tuple[str, ...] | set[str]) -> bool:
    return any(keyword in text for keyword in keywords)


def strip_leading_enumeration(value: str) -> str:
    text = _normalize_output_text(value)
    text = re.sub(r"^[（(]?(?:\d+[.．、])+\d*[）)]?[\s-]*", "", text)
    text = re.sub(r"^[（(]?\d+[）).、\s-]*", "", text)
    text = re.sub(r"^第?[一二三四五六七八九十百千万0-9]+[章节册部分篇节][\s、.．_-]*", "", text)
    return text.strip()


def trim_org_noise_prefix(value: str) -> str:
    text = strip_leading_enumeration(value)
    previous = None
    while text and text != previous:
        previous = text
        text = ORG_NOISE_PREFIX_RE.sub("", text).strip(" ：:-_，。；;|")
    return text


def has_org_noise(value: str) -> bool:
    text = _normalize_output_text(value)
    return contains_any(text, ORG_NOISE_KEYWORDS)


def is_generic_context_segment(value: str) -> bool:
    text = strip_leading_enumeration(Path(value).stem if Path(value).suffix else value)
    if not text:
        return True
    if text in GENERIC_SEGMENTS or text in GENERIC_CONTEXT_SEGMENTS:
        return True
    if any(segment in text for segment in GENERIC_CONTEXT_SEGMENTS) and not has_strong_project_identity(text):
        return True
    if (contains_any(text, REVIEW_HINTS) or contains_any(text, SUPPORT_DOC_HINTS)) and not text.endswith("项目"):
        return True
    if text.startswith("第") and len(text) <= 8:
        return True
    if normalize_project_stage(text) and not has_project_hint(text):
        return True
    return bool(re.fullmatch(r"[A-Za-z]\d+(?:[-(（]\d+[)）])?", text))


def has_strong_project_identity(value: str) -> bool:
    text = strip_leading_enumeration(_normalize_output_text(value))
    if not text:
        return False
    if contains_any(text, PROJECT_NAME_PATH_NOISE_HINTS):
        return False
    return contains_any(text, PROJECT_NAME_STRONG_HINTS) or (has_project_hint(text) and len(text) >= 6)


def should_keep_path_project_candidate(value: str) -> bool:
    text = clean_project_name(value)
    if not text:
        return False
    if contains_any(text, PROJECT_NAME_PATH_NOISE_HINTS):
        return False
    if has_strong_project_identity(text):
        return True
    return has_project_hint(text) and len(text) >= 6 and not re.search(r"\d{4,}", text)


def has_project_hint(value: str) -> bool:
    return any(keyword in _normalize_output_text(value) for keyword in PROJECT_NAME_HINTS)


def is_anchor_document(path_text: str) -> bool:
    normalized = normalize_path_text(path_text)
    return any(keyword in normalized for keyword in CONTEXT_NEUTRAL_HINTS)


def is_attachment_like(path_text: str) -> bool:
    normalized = normalize_path_text(path_text)
    if any(keyword in normalized for keyword in ATTACHMENT_HINTS):
        return True
    return any(keyword in normalized for keyword in SUPPORT_DOC_HINTS) and not is_anchor_document(normalized)


def is_main_project_doc(path_text: str, preferred_stage: str = "") -> bool:
    normalized = normalize_path_text(path_text)
    suffix = Path(normalized).suffix.lower()
    if suffix not in {".doc", ".docx", ".pdf"}:
        return False
    if is_anchor_document(normalized) or is_attachment_like(normalized):
        return False
    stage = normalize_project_stage(normalized)
    if preferred_stage and stage and stage != preferred_stage:
        return False
    stem = Path(normalized).stem
    return has_project_hint(stem) or bool(stage)


def primary_document_weight(item: dict[str, str]) -> int:
    logical_path = item["fileFetchPath"]
    preferred_stage = normalize_project_stage(" ".join(item[field] for field in INPUT_FIELDS))
    if is_anchor_document(logical_path):
        return 4
    if is_main_project_doc(logical_path, preferred_stage):
        return 3
    if is_attachment_like(logical_path):
        return 1
    return 2


def has_negative_unit_context(value: str) -> bool:
    text = _normalize_output_text(value)
    return any(keyword in text for keyword in UNIT_NEGATIVE_CONTEXT_HINTS)


def is_org_candidate(value: str) -> bool:
    text = trim_org_noise_prefix(value)
    if len(text) < 3 or len(text) > 40:
        return False
    if text in GENERIC_SEGMENTS:
        return False
    if has_org_noise(text):
        return False
    if any(marker in text for marker in ORG_REJECT_HINTS):
        return False
    if any(text.endswith(suffix) for suffix in ORG_FACILITY_SUFFIXES):
        return False
    if re.search(r"(项目|工程|建议书|可行性研究|初步设计|竣工验收|目录|封面|标签|概算|正文)$", text):
        return False
    if "|" in text or len(re.findall(r"\d", text)) >= 6:
        return False
    return any(marker in text for marker in ORG_MARKERS)


def clean_org_name(value: str) -> str:
    text = trim_org_noise_prefix(value)
    text = re.sub(r"^关于", "", text)
    text = re.sub(r"^(编制单位|设计单位|建设单位|项目单位|项目法人|业主单位|承担单位|编制机构)[:：]?", "", text)
    text = re.sub(r"(总经理|项目负责人|法定代表人|主管副总经理|主管副所长|院长|董事长).*$", "", text)
    text = re.sub(r"[（(]设计证书.*$", "", text)
    matches = ORG_FRAGMENT_RE.findall(text)
    split_pattern = r"[，。；;:：]|(?:优先采用|采用|不属于|属于|说明|补充说明|文字要求|文本中的位置|要求|增加|实现|推动|开展|通过|由于|其中|由|来自|排至|输送|以及)"
    for fragment in re.split(split_pattern, text):
        fragment = trim_org_noise_prefix(fragment)
        if fragment:
            matches.append(fragment)
    for marker in sorted(ORG_LEGAL_SUFFIXES, key=len, reverse=True):
        for fragment in re.findall(rf"([A-Za-z0-9\u4e00-\u9fff（）()·\-]{{2,60}}?{re.escape(marker)})", text):
            matches.append(fragment)
    matches = [trim_org_noise_prefix(item.strip(" ：:-_，。；;|")) for item in matches if item.strip()]
    matches = list(dict.fromkeys(item for item in matches if item))
    filtered = [item for item in matches if not any(hint in item for hint in ORG_SENTENCE_NOISE_HINTS) and not has_org_noise(item)]
    if filtered:
        matches = filtered
    if matches:
        candidates = [item for item in matches if is_org_candidate(item)]
        if candidates:
            matches = candidates
        matches.sort(key=lambda item: (org_candidate_rank(item), len(item), item), reverse=True)
        text = matches[0]
    return trim_org_noise_prefix(text).strip(" ：:-_，。；;|")


def org_candidate_rank(value: str) -> int:
    text = _normalize_output_text(value)
    if not text:
        return -999
    rank = 0
    if has_org_noise(text):
        rank -= 80
    if any(text.endswith(suffix) for suffix in ORG_FACILITY_SUFFIXES):
        rank -= 60
    if text.endswith(("有限责任公司", "股份有限公司", "有限公司")):
        rank += 40
    elif text.endswith("集团公司"):
        rank += 32
    elif text.endswith(("设计研究院", "研究院", "研究所", "设计院")):
        rank += 24
    elif text.endswith(("公司", "厂", "所", "中心", "学校", "学院", "大学")):
        rank += 16
    elif text.endswith("集团"):
        rank += 12
    if re.match(r"^\d", text):
        rank -= 15
    rank += min(20, len(text))
    return rank


def is_org_alias(shorter: str, longer: str) -> bool:
    short_text = _normalize_output_text(shorter)
    long_text = _normalize_output_text(longer)
    short_key = normalize_for_compare(short_text)
    long_key = normalize_for_compare(long_text)
    if not short_key or not long_key or short_key == long_key or len(short_key) >= len(long_key):
        return False
    if short_key not in long_key and not long_key.startswith(short_key) and not long_key.endswith(short_key):
        return False
    if org_candidate_rank(long_text) <= org_candidate_rank(short_text):
        return False
    if long_text.endswith(ORG_LEGAL_SUFFIXES):
        return True
    return short_text.endswith(("集团", "公司", "厂", "所", "院")) and len(long_key) - len(short_key) <= 10


def merge_org_alias_counter(counter: Counter[str]) -> Counter[str]:
    merged = Counter(counter)
    names = sorted(merged.keys(), key=lambda item: (org_candidate_rank(item), len(item), item), reverse=True)
    for longer in names:
        if merged.get(longer, 0) <= 0:
            continue
        for shorter in list(merged.keys()):
            if shorter == longer or merged.get(shorter, 0) <= 0:
                continue
            if not is_org_alias(shorter, longer):
                continue
            rank_gap = org_candidate_rank(longer) - org_candidate_rank(shorter)
            transfer = merged[shorter] if rank_gap >= 12 else max(1, merged[shorter] // 2)
            merged[longer] += transfer
            merged[shorter] -= transfer
            if merged[shorter] <= 0:
                merged.pop(shorter, None)
    return merged


def looks_like_noise_title(value: str) -> bool:
    text = strip_leading_enumeration(value)
    if not text:
        return True
    if len(text) < 4:
        return True
    if len(text) > 90:
        return True
    if any(keyword in text for keyword in PROJECT_NAME_REJECT_HINTS):
        return True
    if contains_any(text, PROJECT_NAME_HEADING_HINTS):
        return True
    if re.search(r"[.．·]{4,}", text):
        return True
    if re.fullmatch(r"[.．\d\s]*(?:工程|项目)", text):
        return True
    if re.fullmatch(r"第?[一二三四五六七八九十0-9]+[章节册部分篇节].*", text):
        return True
    if re.fullmatch(r"[0-9.\-年月日Vv（）() ]+", text):
        return True
    if "|" in text:
        return True
    if contains_any(text, PROJECT_NAME_PATH_NOISE_HINTS) and not has_strong_project_identity(text):
        return True
    return text in GENERIC_SEGMENTS


def clean_project_name(value: str, unit: str = "") -> str:
    text = strip_leading_enumeration(value)
    if not text:
        return ""
    if any(keyword in text for keyword in ("项目负责人", "项目单位法人", "设计依据", "编制依据")):
        return ""
    if contains_any(text, PROJECT_NAME_HEADING_HINTS):
        return ""
    if "|" in text:
        segments = [segment.strip() for segment in text.split("|") if segment.strip()]
        if not segments:
            return ""
        preferred = [segment for segment in segments if any(keyword in segment for keyword in PROJECT_NAME_HINTS)]
        text = preferred[0] if preferred else segments[0]
    label_match = re.match(r"^(项目名称|工程名称|建设项目名称|项目名|工程名)[:：]\s*(.+)$", text)
    if label_match:
        text = label_match.group(2).strip()
    code_trimmed = re.sub(r"^[A-Za-z]?\d{2,4}(?:[-(（]\d{1,4}[-)）])?", "", text).strip(" ：:-_")
    if code_trimmed and code_trimmed != text and (has_project_hint(code_trimmed) or normalize_project_stage(code_trimmed)):
        text = code_trimmed
    text = re.sub(r"\.[A-Za-z0-9]+$", "", text)
    text = re.sub(r"^第?[一二三四五六七八九十0-9]+[章节册部分篇节][\s、.．_-]*", "", text)
    text = re.sub(r"^(附件|附表|附图)\d*[:：]?", "", text)
    text = re.sub(r"^(封面|目录|标签)[:：]?", "", text)
    text = re.sub(r"[（(][Vv][0-9.]+[）)]", "", text)
    text = re.sub(r"[（(]共?\d+册[）)]", "", text)
    text = re.sub(r"[（(]全一册[）)]", "", text)
    text = re.sub(r"^\d{4}年", "", text)
    text = re.sub(r"\b20\d{2}([.\-年]\d{1,2}){0,2}(日)?\b", "", text)
    text = text.replace("可可研", "可研")
    text = re.sub(r"\b(?:DDD|ddd)\b", "", text)
    text = re.sub(r"(?:审后修改|修改稿|送审稿|蓝图汇总|意见回复|专家意见回复|反馈表|评审意见|对比表|设备表|明细表|分析表)$", "", text)
    for _ in range(3):
        previous = text
        for pattern in PROJECT_SUFFIX_PATTERNS + PROJECT_TRAILING_NOISE_PATTERNS:
            text = re.sub(pattern, "", text)
        text = text.strip(" ：:-_")
        if text == previous:
            break
    text = text.strip(" ：:-_")
    if unit:
        unit_key = normalize_for_compare(unit)
        text_key = normalize_for_compare(text)
        if unit_key and text_key.startswith(unit_key):
            stripped = text[len(unit):].strip()
            if stripped:
                text = stripped
    for _ in range(2):
        match = re.match(r"(.{2,50}?(?:有限公司|有限责任公司|股份有限公司|集团公司|集团|研究院|研究所|设计院|设计研究院|中心|学校|学院|大学|公司|厂|所))(.+)", text)
        if not match:
            break
        org_part = clean_org_name(match.group(1))
        rest = match.group(2).strip(" ：:-_，。；;|")
        if is_org_candidate(org_part) and any(keyword in rest for keyword in PROJECT_NAME_HINTS):
            text = rest
            continue
        break
    text = re.sub(r"\s{2,}", " ", text).strip().strip(" ：:-_，。；;|")
    text = re.sub(r"^[及和]\s*", "", text)
    text = re.sub(r"^年", "", text)
    text = re.sub(r"(?:归档|上报|发文|评审|评估|审后修改|修改稿|送审稿|蓝图汇总)$", "", text).strip(" ：:-_，。；;|")
    if any(keyword in text for keyword in PROJECT_NAME_REJECT_HINTS):
        return ""
    if contains_any(text, PROJECT_NAME_PATH_NOISE_HINTS) and not has_strong_project_identity(text):
        return ""
    if re.search(r"[.．·]{4,}", text):
        return ""
    if re.fullmatch(r"[.．\d\s]*(?:工程|项目)", text):
        return ""
    if len(text) < 4:
        return ""
    if is_org_candidate(text) and not any(keyword in text for keyword in PROJECT_NAME_HINTS):
        return ""
    return "" if looks_like_noise_title(text) else text


def normalize_project_stage(value: str) -> str:
    text = _normalize_output_text(value)
    for stage, keywords in STAGE_KEYWORDS.items():
        if text == stage or any(keyword in text for keyword in keywords):
            return stage
    return ""


def normalize_project_type(value: str) -> str:
    text = _normalize_output_text(value)
    if text in PROJECT_TYPE_VALUES:
        return text
    for project_type, keywords in PROJECT_TYPE_KEYWORDS.items():
        if any(keyword in text for keyword in keywords):
            return project_type
    return ""


def scale_counter(counter: Counter[str], factor: int) -> Counter[str]:
    return Counter({key: value * factor for key, value in counter.items()})


def scale_signal(signal: DocSignals, factor: int) -> DocSignals:
    if factor <= 1:
        return signal
    return DocSignals(
        owner_units=scale_counter(signal.owner_units, factor),
        support_units=scale_counter(signal.support_units, factor),
        explicit_project_names=scale_counter(signal.explicit_project_names, factor),
        general_project_names=scale_counter(signal.general_project_names, factor),
        stage_scores=scale_counter(signal.stage_scores, factor),
        type_scores=scale_counter(signal.type_scores, factor),
    )


def dampen_counter(counter: Counter[str], divisor: int) -> Counter[str]:
    if divisor <= 1:
        return Counter(counter)
    return Counter({key: max(1, value // divisor) for key, value in counter.items() if value > 0})


def consolidate_counter(counter: Counter[str], cleaner) -> Counter[str]:
    merged: Counter[str] = Counter()
    for raw_value, score in counter.items():
        cleaned = cleaner(raw_value)
        if cleaned:
            merged[cleaned] += score
    return merged


def boost_specific_counter(counter: Counter[str]) -> Counter[str]:
    boosted = Counter(counter)
    items = list(counter.items())
    for item, score in items:
        for other, other_score in items:
            if item == other:
                continue
            if item in other and len(other) > len(item):
                boosted[other] += max(1, min(score, other_score) // 4)
    return boosted


def choose_best_counter_entry(counter: Counter[str], ranker=None) -> str:
    if not counter:
        return ""
    best_text = ""
    best_key: tuple[int, int, int, str] | None = None
    for text, score in counter.items():
        rank = ranker(text) if ranker is not None else 0
        key = (score, rank, len(text), text)
        if best_key is None or key > best_key:
            best_text = text
            best_key = key
    return best_text


def choose_specific_counter_entry(counter: Counter[str], ranker, min_score_ratio: float = 0.6) -> str:
    if not counter:
        return ""
    best_score = max(counter.values())
    threshold = max(1, int(best_score * min_score_ratio))
    candidates = [text for text, score in counter.items() if score >= threshold]
    candidates.sort(key=lambda text: (ranker(text), counter[text], len(text), text), reverse=True)
    return candidates[0] if candidates else ""


def project_name_candidate_rank(value: str) -> int:
    text = _normalize_output_text(value)
    rank = 0
    if contains_any(text, PROJECT_NAME_STRONG_HINTS):
        rank += 35
    if has_project_hint(text):
        rank += 15
    if any(keyword in text for keyword in SUPPORT_DOC_HINTS):
        rank -= 25
    if contains_any(text, PROJECT_NAME_PATH_NOISE_HINTS):
        rank -= 40
    if re.search(r"\d{6,}", text):
        rank -= 10
    if not contains_any(text, PROJECT_NAME_STRONG_HINTS):
        rank -= 12
    rank += min(12, len(text))
    return rank


def select_batch_unit(owner_units: Counter[str], support_units: Counter[str]) -> str:
    preferred_owner = choose_specific_counter_entry(owner_units, org_candidate_rank, min_score_ratio=0.55) if owner_units else ""
    if preferred_owner and owner_units.get(preferred_owner, 0) >= 8:
        return preferred_owner
    unit_candidates = Counter(owner_units)
    unit_candidates.update(dampen_counter(support_units, 2))
    unit_candidates = merge_org_alias_counter(unit_candidates)
    preferred_any = choose_specific_counter_entry(unit_candidates, org_candidate_rank, min_score_ratio=0.6)
    return preferred_any or choose_best_counter_entry(unit_candidates, org_candidate_rank) or choose_best_counter_entry(support_units, org_candidate_rank)


def should_prefer_fallback_project_name(current_name: str, fallback_name: str) -> bool:
    current = clean_project_name(current_name)
    fallback = clean_project_name(fallback_name)
    if not fallback:
        return False
    if not current:
        return True
    current_strong = has_strong_project_identity(current)
    fallback_strong = has_strong_project_identity(fallback)
    if current_strong and not fallback_strong:
        return False
    if fallback_strong and not current_strong:
        return True
    current_rank = project_name_candidate_rank(current)
    fallback_rank = project_name_candidate_rank(fallback)
    if current in fallback and len(fallback) > len(current):
        return fallback_rank >= current_rank
    return fallback_rank > current_rank + 8


def payload_name_support_rank(value: str, payload: dict[str, list[str]]) -> int:
    candidate_key = normalize_for_compare(value)
    if not candidate_key:
        return -999
    rank = 0
    matched_fields: set[str] = set()
    matched_items = 0
    for index in range(payload_batch_size(payload)):
        item_score = 0
        for field, base in (("zipName", 12), ("folderName", 10), ("fileName", 6), ("fileFetchPath", 3)):
            raw = normalize_path_text(payload[field][index])
            for segment in [part for part in raw.split("/") if part]:
                probe = Path(segment).stem if segment.endswith(tuple(SUPPORTED_EXTENSIONS)) else segment
                probe_variants = iter_project_name_variants(probe) or [clean_project_name(probe), strip_leading_enumeration(probe)]
                for variant in probe_variants:
                    probe_key = normalize_for_compare(variant)
                    if not probe_key:
                        continue
                    if candidate_key in probe_key or probe_key in candidate_key:
                        item_score = max(item_score, base)
                        matched_fields.add(field)
                        break
        if item_score:
            matched_items += 1
            rank += item_score
    rank += len(matched_fields) * 4
    rank += matched_items * 3
    return rank


def payload_stage_scores(payload: dict[str, list[str]]) -> Counter[str]:
    scores = Counter()
    for index in range(payload_batch_size(payload)):
        for field, base in (("zipName", 14), ("folderName", 12), ("fileName", 10), ("fileFetchPath", 4)):
            raw = normalize_path_text(payload[field][index])
            for segment in [part for part in raw.split("/") if part] or [raw]:
                stage = normalize_project_stage(Path(segment).stem if segment.endswith(tuple(SUPPORTED_EXTENSIONS)) else segment)
                if stage:
                    scores[stage] += base
    return scores


def summarize_signals(signals: list[DocSignals]) -> DocSignals:
    merged = DocSignals()
    for signal in signals:
        merged.owner_units.update(signal.owner_units)
        merged.support_units.update(signal.support_units)
        merged.explicit_project_names.update(signal.explicit_project_names)
        merged.general_project_names.update(signal.general_project_names)
        merged.stage_scores.update(signal.stage_scores)
        merged.type_scores.update(signal.type_scores)
    return merged


def score_context_path(path_text: str, preferred_stage: str = "") -> int:
    normalized = normalize_path_text(path_text)
    stage = normalize_project_stage(normalized)
    if preferred_stage and stage and stage != preferred_stage and not is_anchor_document(normalized):
        return -1
    score = 0
    for index, keyword in enumerate(CONTEXT_DOC_PRIORITY):
        if keyword in normalized:
            score += 100 - index * 8
    if is_anchor_document(normalized):
        score += 80
    if preferred_stage and stage == preferred_stage:
        score += 35
    if is_main_project_doc(normalized, preferred_stage):
        score += 45
    if is_attachment_like(normalized):
        score -= 35
    if any(keyword in normalized for keyword in REVIEW_HINTS):
        score -= 20
    if has_project_hint(Path(normalized).stem):
        score += 12
    if Path(normalized).suffix.lower() in {".doc", ".docx", ".pdf"}:
        score += 10
    return score


def is_supported_extension(path_text: str) -> bool:
    return Path(path_text).suffix.lower() in SUPPORTED_EXTENSIONS


def payload_project_root(payload: dict[str, list[str]]) -> str:
    paths = [normalize_path_text(path).strip("/") for path in payload.get("fileFetchPath", []) if str(path or "").strip()]
    if not paths:
        return ""
    batch_size = payload_batch_size(payload)
    split_paths = [[part for part in path.split("/") if part] for path in paths]
    common_parts: list[str] = []
    for index in range(min(len(parts) for parts in split_paths)):
        probe = split_paths[0][index]
        if all(parts[index] == probe for parts in split_paths):
            common_parts.append(probe)
        else:
            break
    if common_parts and Path(common_parts[-1]).suffix.lower() in SUPPORTED_EXTENSIONS:
        common_parts.pop()
    while common_parts and is_generic_context_segment(common_parts[-1]):
        common_parts.pop()
    while batch_size == 1 and len(common_parts) >= 2 and common_parts[-1] != common_parts[-2]:
        tail = strip_leading_enumeration(common_parts[-1])
        if contains_any(tail, PROJECT_NAME_STRONG_HINTS) or normalize_project_stage(tail) or tail.endswith("归档"):
            break
        common_parts.pop()
    if common_parts:
        return "/".join(common_parts)
    first_parts = split_paths[0]
    if first_parts and Path(first_parts[-1]).suffix.lower() in SUPPORTED_EXTENSIONS:
        first_parts = first_parts[:-1]
    while first_parts and is_generic_context_segment(first_parts[-1]):
        first_parts = first_parts[:-1]
    return "/".join(first_parts)


def build_context_item(zip_name: str, logical_path: str) -> dict[str, str]:
    normalized = normalize_path_text(logical_path).strip("/")
    path = Path(normalized)
    return {
        "zipName": zip_name,
        "folderName": path.parent.name,
        "fileName": path.name,
        "fileFetchPath": normalized,
    }


def weighted_payload_name_candidates(payload: dict[str, list[str]]) -> Counter[str]:
    counter: Counter[str] = Counter()
    for index in range(payload_batch_size(payload)):
        item = {field: payload[field][index] for field in INPUT_FIELDS}
        for value, weight in (
            (item["zipName"], 10),
            (item["folderName"], 8),
            (item["fileName"], 5),
            (item["fileFetchPath"], 2),
        ):
            for segment in normalize_path_text(value).split("/"):
                cleaned = segment.strip()
                if not cleaned or cleaned in GENERIC_SEGMENTS:
                    continue
                if cleaned.endswith(tuple(SUPPORTED_EXTENSIONS)):
                    cleaned = Path(cleaned).stem
                if re.fullmatch(r"[A-Za-z]\d+[\w().\-]*", cleaned):
                    continue
                for project_name in iter_project_name_variants(cleaned):
                    if not should_keep_path_project_candidate(project_name):
                        continue
                    candidate_weight = weight + max(0, project_name_candidate_rank(project_name) // 10)
                    if is_attachment_like(cleaned):
                        candidate_weight = max(1, candidate_weight - 4)
                    counter[project_name] += max(1, candidate_weight)
    return counter


def build_logger(config: dict[str, Any]) -> logging.Logger:
    logger = logging.getLogger("zch_extractor")
    logger.setLevel(getattr(logging, str(config["logging"].get("level", "INFO")).upper(), logging.INFO))
    logger.propagate = False
    if not logger.handlers:
        handler = logging.StreamHandler(sys.stdout)
        handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
        logger.addHandler(handler)
    return logger


def parse_json_object(content: str) -> Any:
    if not content:
        return None
    text = content.strip()
    try:
        return json.loads(text)
    except Exception:
        pass
    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        try:
            return json.loads(text[start : end + 1])
        except Exception:
            return None
    return None


def render_prompt_template(template: str, **values: str) -> str:
    rendered = template
    for key, value in values.items():
        rendered = rendered.replace("{" + key + "}", value)
    return rendered


def load_runtime_config(config_path: str | Path) -> dict[str, Any]:
    path = Path(config_path).resolve()
    if not path.exists():
        raise FileNotFoundError(f"config.json 不存在: {path}")
    raw = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(raw, dict):
        raise ValueError("config.json 顶层必须是对象")
    llm = raw.get("llm", {}) if isinstance(raw.get("llm"), dict) else {}
    extraction = raw.get("extraction", {}) if isinstance(raw.get("extraction"), dict) else {}
    api_model = raw.get("api_model", {}) if isinstance(raw.get("api_model"), dict) else {}
    local_model = raw.get("local_model", {}) if isinstance(raw.get("local_model"), dict) else {}
    paths = raw.get("paths", {}) if isinstance(raw.get("paths"), dict) else {}
    logging_cfg = raw.get("logging", {}) if isinstance(raw.get("logging"), dict) else {}
    output = raw.get("output", {}) if isinstance(raw.get("output"), dict) else {}
    ocr = raw.get("ocr", {}) if isinstance(raw.get("ocr"), dict) else {}
    minio = raw.get("minio", {}) if isinstance(raw.get("minio"), dict) else {}
    if not api_model and llm:
        api_model = {
            "enabled": llm.get("enabled", True),
            "base_url": llm.get("base_url", ""),
            "api_key": llm.get("api_key", ""),
            "model": llm.get("model", ""),
            "temperature": llm.get("temperature", 0.0),
            "max_tokens": llm.get("max_tokens", 512),
            "timeout_seconds": llm.get("timeout_seconds", 30),
        }
    return {
        "test_mode": _as_bool(raw.get("test_mode", True), True),
        "minio": {
            "endpoint": str(minio.get("endpoint", "") or "").strip(),
            "sdk_endpoint": str(minio.get("sdk_endpoint", "") or "").strip(),
            "access_key": str(minio.get("access_key", "") or "").strip(),
            "secret_key": str(minio.get("secret_key", "") or "").strip(),
            "bucket": str(minio.get("bucket", "") or "").strip(),
            "source_prefix": str(minio.get("source_prefix", "") or "").strip(),
            "result_prefix": str(minio.get("result_prefix", "") or "").strip(),
            "secure": minio.get("secure"),
            "allow_local_fallback": _as_bool(minio.get("allow_local_fallback", True), True),
            "prefer_remote_in_test": _as_bool(minio.get("prefer_remote_in_test", False), False),
            "strict_backup_in_production": _as_bool(minio.get("strict_backup_in_production", False), False),
        },
        "api_model": {
            "enabled": _as_bool(api_model.get("enabled", True), True),
            "base_url": str(api_model.get("base_url", "") or "").strip(),
            "api_key": str(api_model.get("api_key", "") or "").strip(),
            "model": str(api_model.get("model", "") or "").strip(),
            "temperature": float(api_model.get("temperature", 0.0) or 0.0),
            "max_tokens": _as_int(api_model.get("max_tokens", 512), 512),
            "timeout_seconds": _as_int(api_model.get("timeout_seconds", 30), 30),
        },
        "local_model": {
            "enabled": _as_bool(local_model.get("enabled", True), True),
            "base_url": str(local_model.get("base_url", "") or "").strip(),
            "api_key": str(local_model.get("api_key", "") or "").strip(),
            "model": str(local_model.get("model", "") or "").strip(),
            "temperature": float(local_model.get("temperature", 0.0) or 0.0),
            "max_tokens": _as_int(local_model.get("max_tokens", 512), 512),
            "timeout_seconds": _as_int(local_model.get("timeout_seconds", 30), 30),
        },
        "ocr": {
            "enabled": _as_bool(ocr.get("enabled", True), True),
            "backend": str(ocr.get("backend", "rapidocr") or "rapidocr"),
            "dpi": _as_int(ocr.get("dpi", 300), 300),
            "max_pdf_pages": _as_int(ocr.get("max_pdf_pages", extraction.get("max_pdf_pages", 8)), 8),
            "max_excel_rows": _as_int(ocr.get("max_excel_rows", 30), 30),
            "max_excel_cols": _as_int(ocr.get("max_excel_cols", 12), 12),
            "max_front_chars": _as_int(ocr.get("max_front_chars", 4000), 4000),
            "title_scan_lines": _as_int(ocr.get("title_scan_lines", extraction.get("title_scan_lines", 80)), 80),
            "only_when_text_empty": _as_bool(ocr.get("only_when_text_empty", True), True),
            "min_chars_to_skip_ocr": _as_int(ocr.get("min_chars_to_skip_ocr", 80), 80),
        },
        "paths": {
            "code_root": str(Path(paths.get("code_root", PROJECT_ROOT)).resolve()),
            "data_root": str(Path(paths.get("data_root", raw.get("local_source_root", PROJECT_ROOT))).resolve()),
            "local_test_root": str(Path(paths.get("local_test_root", Path(paths.get("data_root", raw.get("local_source_root", PROJECT_ROOT))) / "测试项目")).resolve()),
            "prompt_path": str(Path(paths.get("prompt_path", path.with_name("prompt.json"))).resolve()),
            "filejsonrst_path": str(Path(paths.get("filejsonrst_path", path.with_name("filejsonrst.json"))).resolve()),
        },
        "logging": {"level": str(logging_cfg.get("level", "INFO"))},
        "output": {
            "tag_filename": str(output.get("tag_filename", raw.get("output_name", "tag.json")) or "tag.json"),
            "persist_single_mode": _as_bool(output.get("persist_single_mode", False), False),
            "indent": _as_int(output.get("indent", 2), 2),
            "ensure_ascii": _as_bool(output.get("ensure_ascii", False), False),
        },
    }


def choose_best_text(counter: Counter[str]) -> str:
    if not counter:
        return ""
    best_score = max(counter.values())
    candidates = [item for item, score in counter.items() if score == best_score]
    candidates.sort(key=lambda item: (any(marker in item for marker in ("有限责任公司", "股份有限公司", "研究院", "研究所", "集团")), len(item), item), reverse=True)
    return candidates[0]


def decode_escaped_unicode_name(name: str) -> str:
    return ESCAPED_UNICODE_RE.sub(lambda match: chr(int(match.group(1), 16)), name or "")


def encode_path_to_escaped_unicode(raw_path: str) -> str:
    normalized = str(raw_path or "").replace("\\", "/")
    parts: list[str] = []
    for part in normalized.split("/"):
        if not part:
            continue
        encoded = []
        for char in part:
            code = ord(char)
            encoded.append(f"#U{code:04x}" if code > 127 else char)
        parts.append("".join(encoded))
    return "/".join(parts)


def path_candidates_for_lookup(raw_path: str) -> list[str]:
    normalized = str(raw_path or "").strip().replace("\\", "/")
    candidates: list[str] = []
    for value in (normalized, decode_escaped_unicode_name(normalized), encode_path_to_escaped_unicode(normalized)):
        cleaned = value.strip("/")
        if cleaned and cleaned not in candidates:
            candidates.append(cleaned)
    return candidates


def resolve_reference_path(raw_path: str, base_dir: Path, fallback_dir: Path) -> Path:
    candidate = Path(raw_path)
    if candidate.is_absolute():
        return candidate.resolve()
    first = (base_dir / candidate).resolve()
    return first if first.exists() else (fallback_dir / candidate).resolve()

def normalize_payload(payload: dict[str, Any]) -> dict[str, list[str]]:
    normalized: dict[str, list[str]] = {}
    for key in INPUT_FIELDS:
        value = payload.get(key, [])
        normalized[key] = ["" if item is None else str(item) for item in value] if isinstance(value, list) else []
    return normalized


def payload_batch_size(payload: dict[str, list[str]]) -> int:
    return max((len(payload.get(key, [])) for key in INPUT_FIELDS), default=0)


def validate_precheck(payload: dict[str, list[str]]) -> list[str]:
    lengths = [len(payload.get(key, [])) for key in INPUT_FIELDS]
    if not lengths or max(lengths) == 0:
        return ["输入批次为空"]
    if len(set(lengths)) != 1:
        return ["zipName、folderName、fileName、fileFetchPath 列表长度不一致"]
    reasons: list[str] = []
    for index in range(lengths[0]):
        path_text = normalize_path_text(payload["fileFetchPath"][index])
        if payload["fileName"][index] not in path_text:
            reasons.append(f"第 {index + 1} 条 fileName 未严格包含于 fileFetchPath")
        if payload["folderName"][index] not in path_text:
            reasons.append(f"第 {index + 1} 条 folderName 未严格包含于 fileFetchPath")
    return reasons


def build_fail_output(size: int) -> dict[str, Any]:
    return {
        "unit": ["" for _ in range(size)],
        "product": ["" for _ in range(size)],
        "project_type": ["" for _ in range(size)],
        "project_name": ["" for _ in range(size)],
        "project_stage": ["" for _ in range(size)],
        "status": "fail",
    }


def build_success_output(size: int, fields: dict[str, str]) -> dict[str, Any]:
    return {
        "unit": [fields.get("unit", "") for _ in range(size)],
        "product": ["N/A" for _ in range(size)],
        "project_type": [fields.get("project_type", "") for _ in range(size)],
        "project_name": [fields.get("project_name", "") for _ in range(size)],
        "project_stage": [fields.get("project_stage", "") for _ in range(size)],
        "status": "success",
    }


def validate_postcheck(payload: dict[str, list[str]], output: dict[str, Any]) -> list[str]:
    expected = payload_batch_size(payload)
    reasons: list[str] = []
    for field in OUTPUT_FIELDS:
        value = output.get(field)
        if not isinstance(value, list) or len(value) != expected:
            reasons.append(f"输出字段 {field} 长度与输入批次长度不一致")
    if output.get("status") not in {"success", "fail"}:
        reasons.append("status 非法")
    return reasons


def build_minio_object_candidates(source_prefix: str, file_fetch_path: str) -> list[str]:
    normalized = normalize_path_text(file_fetch_path).lstrip("/")
    prefix = normalize_path_text(source_prefix).strip("/")
    candidates: list[str] = []
    if prefix:
        candidates.append(normalized if normalized.startswith(f"{prefix}/") else f"{prefix}/{normalized}")
    candidates.append(normalized)
    output: list[str] = []
    seen: set[str] = set()
    for candidate in candidates:
        if candidate and candidate not in seen:
            seen.add(candidate)
            output.append(candidate)
    return output


def _candidate_minio_endpoints(minio_cfg: dict[str, Any]) -> list[str]:
    endpoints: list[str] = []
    for key in ("sdk_endpoint", "endpoint"):
        value = str(minio_cfg.get(key, "") or "").strip()
        if value:
            endpoints.append(value)
            parsed = urllib.parse.urlparse(value if "://" in value else f"http://{value}")
            if parsed.port == 9001 and parsed.hostname:
                endpoints.append(urllib.parse.urlunparse((parsed.scheme or "http", f"{parsed.hostname}:9000", "", "", "", "")))
    output: list[str] = []
    seen: set[str] = set()
    for endpoint in endpoints:
        if endpoint not in seen:
            seen.add(endpoint)
            output.append(endpoint)
    return output


def check_minio(minio_cfg: dict[str, Any], logger: logging.Logger) -> MinioState:
    if Minio is None:
        return MinioState(False, None, "minio 包未安装", "")
    bucket = str(minio_cfg.get("bucket", "") or "").strip()
    if not bucket:
        return MinioState(False, None, "bucket 未配置", "")
    access_key = str(minio_cfg.get("access_key", "") or "").strip()
    secret_key = str(minio_cfg.get("secret_key", "") or "").strip()
    errors: list[str] = []
    for endpoint in _candidate_minio_endpoints(minio_cfg):
        parsed = urllib.parse.urlparse(endpoint if "://" in endpoint else f"http://{endpoint}")
        host = parsed.netloc or parsed.path
        secure = bool(minio_cfg.get("secure")) if minio_cfg.get("secure") is not None else parsed.scheme == "https"
        if not host:
            continue
        try:
            client = Minio(host, access_key=access_key, secret_key=secret_key, secure=secure)
            if client.bucket_exists(bucket):
                logger.info("MinIO 连通成功: %s", endpoint)
                return MinioState(True, client, "", endpoint)
        except Exception as exc:
            errors.append(f"{endpoint}: {exc}")
    return MinioState(False, None, " | ".join(errors) or "未配置可用 MinIO endpoint", "")


def probe_openai_compatible(model_cfg: dict[str, Any]) -> tuple[bool, str]:
    if not _as_bool(model_cfg.get("enabled", True), True):
        return False, "模型开关关闭"
    base_url = str(model_cfg.get("base_url", "") or "").strip()
    if not base_url:
        return False, "base_url 未配置"
    headers = {"Content-Type": "application/json"}
    api_key = str(model_cfg.get("api_key", "") or "").strip()
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    request = urllib.request.Request(url=f"{base_url.rstrip('/')}/models", headers=headers, method="GET")
    try:
        with urllib.request.urlopen(request, timeout=_as_int(model_cfg.get("timeout_seconds", 30), 30)) as response:
            if response.status < 500:
                return True, ""
    except urllib.error.HTTPError as exc:
        if exc.code == 404:
            return True, ""
        return False, f"HTTP {exc.code}"
    except Exception as exc:
        return False, str(exc)
    return False, "未知错误"


def choose_model_provider(config: dict[str, Any], logger: logging.Logger) -> ModelProvider | None:
    api_ok, api_reason = probe_openai_compatible(config["api_model"])
    if api_ok:
        logger.info("公网模型可用")
        return ModelProvider("api_model", config["api_model"])
    logger.warning("公网模型失败: %s", api_reason)
    local_ok, local_reason = probe_openai_compatible(config["local_model"])
    if local_ok:
        logger.info("本地模型可用")
        return ModelProvider("local_model", config["local_model"])
    logger.warning("本地模型失败: %s", local_reason)
    logger.warning("进入无模型模式")
    return None


def call_openai_compatible(provider: ModelProvider, system_prompt: str, user_prompt: str) -> str:
    cfg = provider.config
    headers = {"Content-Type": "application/json"}
    api_key = str(cfg.get("api_key", "") or "").strip()
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    body = json.dumps(
        {
            "model": cfg.get("model") or "",
            "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            "temperature": float(cfg.get("temperature", 0.0) or 0.0),
            "max_tokens": _as_int(cfg.get("max_tokens", 512), 512),
        },
        ensure_ascii=False,
    ).encode("utf-8")
    request = urllib.request.Request(url=f"{str(cfg.get('base_url', '')).rstrip('/')}/chat/completions", data=body, headers=headers, method="POST")
    with urllib.request.urlopen(request, timeout=_as_int(cfg.get("timeout_seconds", 30), 30)) as response:
        payload = json.loads(response.read().decode("utf-8"))
    choices = payload.get("choices", [])
    if not choices:
        raise ValueError("模型返回为空")
    return str(choices[0].get("message", {}).get("content", "") or "")

def split_useful_lines(text: str, limit: int) -> list[str]:
    lines: list[str] = []
    for raw in text.splitlines():
        line = _normalize_output_text(raw)
        if line and line not in lines:
            lines.append(line)
        if len(lines) >= limit:
            break
    return lines


def extract_unit_candidates(lines: list[str], text: str, front_org_to_owner: bool = True) -> tuple[Counter[str], Counter[str]]:
    owner, support = Counter(), Counter()
    for label, weight in UNIT_LABEL_WEIGHTS.items():
        patterns = (rf"{label}[:：]?\s*([^\n]{{2,100}})", rf"{label}\s*\n([^\n]{{2,100}})")
        for pattern in patterns:
            for candidate in re.findall(pattern, text):
                cleaned = clean_org_name(candidate)
                if is_org_candidate(cleaned) and not has_negative_unit_context(candidate):
                    (owner if label in {"建设单位", "项目单位", "项目法人", "业主单位", "建设管理单位", "申报单位", "承担单位"} else support)[cleaned] += weight
    stage_index = next((index for index, line in enumerate(lines[:15]) if normalize_project_stage(line)), -1)
    if stage_index >= 0:
        for index, line in enumerate(lines[: stage_index + 1]):
            if has_negative_unit_context(line):
                continue
            cleaned = clean_org_name(line)
            if is_org_candidate(cleaned):
                target = owner if front_org_to_owner else support
                target[cleaned] += max(2, 12 - (stage_index - index)) if front_org_to_owner else max(2, 6 - (stage_index - index))
        for line in lines[stage_index + 1 : stage_index + 7]:
            if has_negative_unit_context(line):
                continue
            cleaned = clean_org_name(line)
            if is_org_candidate(cleaned):
                support[cleaned] += 4
    else:
        for line in lines[:8]:
            if has_negative_unit_context(line):
                continue
            cleaned = clean_org_name(line)
            if is_org_candidate(cleaned):
                target = owner if front_org_to_owner else support
                target[cleaned] += 5 if front_org_to_owner else 2
    for pattern in (r"(?:项目单位概况|建设单位基本情况|单位基本情况)\s*\n([^\n]{2,100})", r"(?:项目法人|建设单位)[:：]?\s*([^\n]{2,100})"):
        for candidate in re.findall(pattern, text):
            cleaned = clean_org_name(candidate)
            if is_org_candidate(cleaned) and not has_negative_unit_context(candidate):
                owner[cleaned] += 10
    return owner, support


def iter_project_name_variants(value: str, unit: str = "") -> list[str]:
    raw = _normalize_output_text(value)
    if not raw:
        return []
    raw = re.sub(r"\.[A-Za-z0-9]+$", "", raw)
    variants = [raw]
    trimmed_code = re.sub(r"^[A-Za-z]?\d{2,4}(?:[-(（]\d{1,4}[-)）])?", "", raw).strip(" ：:-_")
    if trimmed_code and trimmed_code != raw:
        variants.append(trimmed_code)
    variants.append(re.sub(r"(?:归档|上报|发文|评审|评估|审后修改|修改稿|送审稿|蓝图汇总).*$", "", raw).strip(" ：:-_"))
    variants.append(re.sub(r"(?:项目建议书|建议书|可行性研究报告|可行性研究|可研|初步设计|初设|竣工验收总结报告|竣工验收报告|竣工验收|竣工|专项规划|总体规划|规划).*$", "", raw).strip(" ：:-_"))
    output: list[str] = []
    seen: set[str] = set()
    for candidate in variants:
        cleaned = clean_project_name(candidate, unit)
        if cleaned and cleaned not in seen:
            seen.add(cleaned)
            output.append(cleaned)
    return output


def extract_project_name_candidates(lines: list[str], text: str, path_candidates: list[str]) -> tuple[Counter[str], Counter[str]]:
    explicit, general = Counter(), Counter()
    for pattern in (r"(?:项目名称|工程名称|建设项目名称|项目名称，项目法人|项目名称、项目法人)[:：]?\s*([^\n]{2,120})", r"关于([^\n]{4,120}?(?:项目|工程))(?:可行性研究报告|项目建议书|初步设计|竣工验收报告|的批复)"):
        for candidate in re.findall(pattern, text):
            for cleaned in iter_project_name_variants(candidate):
                explicit[cleaned] += 16
    for match in TITLE_HINT_RE.findall("\n".join(lines[:18])):
        for cleaned in iter_project_name_variants(match):
            general[cleaned] += 10
    stage_index = next((index for index, line in enumerate(lines[:15]) if normalize_project_stage(line)), -1)
    if stage_index > 0:
        for line in lines[max(0, stage_index - 2) : stage_index]:
            for cleaned in iter_project_name_variants(line):
                general[cleaned] += 12
    for line in lines[:12]:
        if looks_like_noise_title(line) or is_org_candidate(line):
            continue
        if any(keyword in line for keyword in ("项目", "工程", "规划", "办法", "条例", "改造", "建设", "车间")):
            for cleaned in iter_project_name_variants(line):
                general[cleaned] += 8
    for candidate in path_candidates:
        for cleaned in iter_project_name_variants(candidate):
            general[cleaned] += 5
    return explicit, general


def score_stage(path_text: str, title_text: str, body_text: str) -> Counter[str]:
    scores = Counter()
    support_like = any(keyword in path_text or keyword in title_text for keyword in STAGE_REFERENCE_HINTS)
    for stage, keywords in STAGE_KEYWORDS.items():
        for keyword in keywords:
            if keyword in path_text:
                scores[stage] += 8
            if keyword in title_text:
                scores[stage] += 10
            if not support_like and keyword in body_text:
                scores[stage] += 1
    return scores


def score_project_type(path_text: str, title_text: str, body_text: str, stage: str) -> Counter[str]:
    scores = Counter()
    if stage:
        scores["固定资产投资"] += 12
    for project_type, keywords in PROJECT_TYPE_KEYWORDS.items():
        for keyword in keywords:
            if keyword in path_text:
                scores[project_type] += 6
            if keyword in title_text:
                scores[project_type] += 8
            if keyword in body_text:
                scores[project_type] += 1
    if scores["政策法规"] and scores["固定资产投资"] and "专家意见" in path_text:
        scores["政策法规"] = max(0, scores["政策法规"] - 8)
    return scores


def informative_path_candidates(item: dict[str, str]) -> list[str]:
    candidates: list[str] = []
    for value in (item["zipName"], item["folderName"], item["fileName"], item["fileFetchPath"]):
        for segment in normalize_path_text(value).split("/"):
            cleaned = segment.strip()
            if not cleaned or cleaned in GENERIC_SEGMENTS:
                continue
            if cleaned.endswith((".doc", ".docx", ".pdf", ".xls", ".xlsx")):
                cleaned = Path(cleaned).stem
            if re.fullmatch(r"[A-Za-z]\d+[\w().\-]*", cleaned):
                continue
            candidates.extend(candidate for candidate in iter_project_name_variants(cleaned) if should_keep_path_project_candidate(candidate))
    output: list[str] = []
    seen: set[str] = set()
    for candidate in candidates:
        if candidate not in seen:
            seen.add(candidate)
            output.append(candidate)
    return output


def analyze_document(item: dict[str, str], extracted: ExtractedDocument, config: dict[str, Any]) -> DocSignals:
    signal = DocSignals()
    lines = split_useful_lines(extracted.front_text, int(config["ocr"].get("title_scan_lines", 80)))
    path_candidates = informative_path_candidates(item)
    front_org_to_owner = is_anchor_document(item["fileFetchPath"]) or is_main_project_doc(item["fileFetchPath"], normalize_project_stage(" ".join(item[field] for field in INPUT_FIELDS)))
    owner_units, support_units = extract_unit_candidates(lines, extracted.text, front_org_to_owner=front_org_to_owner)
    signal.owner_units.update(owner_units)
    signal.support_units.update(support_units)
    explicit_names, general_names = extract_project_name_candidates(lines, extracted.text, path_candidates)
    signal.explicit_project_names.update(explicit_names)
    signal.general_project_names.update(general_names)
    path_text = " ".join(item[field] for field in INPUT_FIELDS)
    title_text = "\n".join(lines[:15])
    stage_body_text = "\n".join(lines[:30]) or extracted.front_text[:2500]
    signal.stage_scores.update(score_stage(path_text, title_text, stage_body_text))
    signal.type_scores.update(score_project_type(path_text, title_text, extracted.text[:6000], normalize_project_stage(choose_best_text(signal.stage_scores))))
    return signal


def aggregate_batch_fields(payload: dict[str, list[str]], primary_signals: list[DocSignals], context_signals: list[DocSignals] | None = None) -> dict[str, str]:
    primary = summarize_signals(primary_signals)
    context = summarize_signals(context_signals or [])
    owner_units = merge_org_alias_counter(consolidate_counter(primary.owner_units, clean_org_name))
    owner_units.update(dampen_counter(merge_org_alias_counter(consolidate_counter(context.owner_units, clean_org_name)), 2))
    owner_units = merge_org_alias_counter(owner_units)
    support_units = merge_org_alias_counter(consolidate_counter(primary.support_units, clean_org_name))
    support_units.update(dampen_counter(merge_org_alias_counter(consolidate_counter(context.support_units, clean_org_name)), 2))
    support_units = merge_org_alias_counter(support_units)
    batch_unit = select_batch_unit(owner_units, support_units)
    stage_scores = Counter(primary.stage_scores)
    stage_scores.update(dampen_counter(context.stage_scores, 2))
    path_stage_scores = payload_stage_scores(payload)
    stage_scores.update(path_stage_scores)
    batch_stage = normalize_project_stage(choose_specific_counter_entry(stage_scores, lambda value: path_stage_scores.get(value, 0), min_score_ratio=0.75))
    if not batch_stage:
        batch_stage = normalize_project_stage(choose_best_counter_entry(stage_scores)) or normalize_project_stage(" ".join(payload.get("folderName", []) + payload.get("fileName", []) + payload.get("zipName", [])))
    type_scores = Counter(primary.type_scores)
    type_scores.update(dampen_counter(context.type_scores, 2))
    batch_type = normalize_project_type(choose_best_counter_entry(type_scores)) or ("固定资产投资" if batch_stage else "其他")
    explicit_names = consolidate_counter(primary.explicit_project_names, lambda value: clean_project_name(value, batch_unit))
    explicit_names.update(dampen_counter(consolidate_counter(context.explicit_project_names, lambda value: clean_project_name(value, batch_unit)), 2))
    general_names = boost_specific_counter(consolidate_counter(primary.general_project_names, lambda value: clean_project_name(value, batch_unit)))
    general_names.update(dampen_counter(boost_specific_counter(consolidate_counter(context.general_project_names, lambda value: clean_project_name(value, batch_unit))), 2))
    fallback = boost_specific_counter(weighted_payload_name_candidates(payload))
    name_candidates = Counter(general_names)
    name_candidates.update(scale_counter(explicit_names, 2))
    name_candidates.update(dampen_counter(fallback, 3))
    name_ranker = lambda value: project_name_candidate_rank(value) + payload_name_support_rank(value, payload) * 2
    batch_name = choose_specific_counter_entry(boost_specific_counter(name_candidates), name_ranker, min_score_ratio=0.7) or choose_best_counter_entry(boost_specific_counter(name_candidates), name_ranker)
    fallback_name = clean_project_name(choose_best_counter_entry(fallback, name_ranker), batch_unit)
    if not batch_name or looks_like_noise_title(batch_name):
        batch_name = fallback_name
    elif fallback_name and should_prefer_fallback_project_name(batch_name, fallback_name):
        batch_name = fallback_name
    payload_preferred_name = choose_specific_counter_entry(
        boost_specific_counter(name_candidates),
        lambda value: payload_name_support_rank(value, payload) + project_name_candidate_rank(value),
        min_score_ratio=0.45,
    )
    if payload_preferred_name:
        current_support = payload_name_support_rank(batch_name, payload)
        preferred_support = payload_name_support_rank(payload_preferred_name, payload)
        if preferred_support >= current_support + 20 and project_name_candidate_rank(payload_preferred_name) >= project_name_candidate_rank(batch_name) - 10:
            batch_name = payload_preferred_name
    return {"unit": batch_unit, "product": "N/A", "project_type": batch_type, "project_name": batch_name, "project_stage": batch_stage}


def _doc_priority(logical_path: str) -> int:
    name = normalize_path_text(logical_path)
    score = 0
    for index, keyword in enumerate(("标签", "封面", "目录", "建议书", "可研", "初设", "竣工", "正文", "报告", "概算")):
        if keyword in name:
            score += 20 - index
    return score

def process_request(payload: dict[str, Any], config_path: str | Path = PROJECT_ROOT / "config.json", persist_result: bool = False, runtime: RuntimeContext | None = None, prompt_override: str = "", override_test_mode: bool | None = None) -> dict[str, Any]:
    created_runtime = runtime is None
    context = runtime or RuntimeContext(config_path, override_test_mode=override_test_mode)
    try:
        normalized = normalize_payload(payload if isinstance(payload, dict) else {})
        batch_size = payload_batch_size(normalized)
        precheck_errors = validate_precheck(normalized)
        if precheck_errors:
            context.logger.error("前置检查失败: %s", " | ".join(precheck_errors))
            return build_fail_output(batch_size)
        documents: list[ExtractedDocument] = []
        for index in range(batch_size):
            item = {field: normalized[field][index] for field in INPUT_FIELDS}
            extracted = context.resolve_document(item)
            documents.append(extracted)
            if extracted.errors:
                context.logger.error("关键文件读取失败: %s -> %s", item["fileFetchPath"], " | ".join(extracted.errors))
                return build_fail_output(batch_size)
        primary_signals = [
            scale_signal(
                analyze_document({field: normalized[field][index] for field in INPUT_FIELDS}, documents[index], context.config),
                primary_document_weight({field: normalized[field][index] for field in INPUT_FIELDS}),
            )
            for index in range(batch_size)
        ]
        context_signals: list[DocSignals] = []
        for item, extracted, weight in context.discover_project_context(normalized):
            context_signals.append(scale_signal(analyze_document(item, extracted, context.config), weight))
            documents.append(extracted)
        fields = aggregate_batch_fields(normalized, primary_signals, context_signals)
        fields = context.maybe_refine_with_llm(normalized, fields, documents, prompt_override=prompt_override)
        output = build_success_output(batch_size, fields)
        postcheck_errors = validate_postcheck(normalized, output)
        if postcheck_errors:
            context.logger.error("后置检查失败: %s", " | ".join(postcheck_errors))
            return build_fail_output(batch_size)
        if persist_result:
            persist_batch_results([output], context, allow_local_write=context.config["test_mode"] or context.minio.available)
        context.logger.info("最终结果 success")
        return output
    except Exception as exc:
        context.logger.error("程序异常，返回 fail: %s", exc)
        return build_fail_output(payload_batch_size(normalize_payload(payload if isinstance(payload, dict) else {})))
    finally:
        if created_runtime:
            context.close()


def persist_batch_results(results: list[dict[str, Any]], runtime: RuntimeContext, allow_local_write: bool) -> None:
    if runtime.minio.available and runtime.minio.client is not None:
        bucket = str(runtime.config["minio"].get("bucket", "")).strip()
        result_prefix = normalize_path_text(str(runtime.config["minio"].get("result_prefix", "") or "").strip()).strip("/")
        tag_name = runtime.config["output"]["tag_filename"]
        object_name = f"{result_prefix}/{tag_name}" if result_prefix else tag_name
        data = json.dumps(results, ensure_ascii=runtime.config["output"]["ensure_ascii"], indent=runtime.config["output"]["indent"]).encode("utf-8")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".json") as tmp:
            tmp.write(data)
            temp_path = Path(tmp.name)
        try:
            runtime.minio.client.fput_object(bucket, object_name, str(temp_path), content_type="application/json")
        finally:
            temp_path.unlink(missing_ok=True)
    elif allow_local_write:
        tag_path = Path(runtime.config["paths"]["code_root"]) / runtime.config["output"]["tag_filename"]
        tag_path.write_text(json.dumps(results, ensure_ascii=runtime.config["output"]["ensure_ascii"], indent=runtime.config["output"]["indent"]), encoding="utf-8")
    else:
        runtime.logger.warning("MinIO 不可用且未显式启用本地写入，仍回退写入本地 tag.json")
        tag_path = Path(runtime.config["paths"]["code_root"]) / runtime.config["output"]["tag_filename"]
        tag_path.write_text(json.dumps(results, ensure_ascii=runtime.config["output"]["ensure_ascii"], indent=runtime.config["output"]["indent"]), encoding="utf-8")


def process_manifest(filejsonrst_path: str | Path, config_path: str | Path = PROJECT_ROOT / "config.json", override_test_mode: bool | None = None) -> list[dict[str, Any]]:
    runtime = RuntimeContext(config_path, override_test_mode=override_test_mode)
    try:
        manifest_path = Path(filejsonrst_path).resolve()
        if not manifest_path.exists():
            runtime.logger.error("filejsonrst.json 不存在: %s", manifest_path)
            return []
        data = json.loads(manifest_path.read_text(encoding="utf-8"))
        if not isinstance(data, list):
            runtime.logger.error("filejsonrst.json 顶层必须为数组")
            return []
        if not runtime.config["test_mode"] and not runtime.minio.available and runtime.config["minio"].get("strict_backup_in_production", False):
            runtime.logger.error("生产模式 MinIO 检查失败，且 strict_backup_in_production=true，批量直接 fail")
            return [build_fail_output(payload_batch_size(normalize_payload(item if isinstance(item, dict) else {}))) for item in data]
        results: list[dict[str, Any]] = []
        for item in data:
            payload = item if isinstance(item, dict) else {}
            prompt_override = str(payload.get("llmPromptConfigPath", "") or "").strip() if isinstance(payload, dict) else ""
            results.append(process_request(payload, config_path=config_path, persist_result=False, runtime=runtime, prompt_override=prompt_override))
        allow_local_write = runtime.config["test_mode"] or runtime.minio.available
        if allow_local_write or runtime.minio.available:
            persist_batch_results(results, runtime, allow_local_write=allow_local_write)
        return results
    finally:
        runtime.close()


def run_payload(payload: dict[str, Any], config_path: str = "./config.json", override_test_mode: bool | None = None) -> dict[str, Any]:
    return process_request(payload, config_path=config_path, persist_result=False, override_test_mode=override_test_mode)


def run_job(config_path: str, override_test_mode: bool | None = None, disable_llm: bool = False) -> dict[str, Any]:
    if disable_llm:
        pass
    results = process_manifest(load_runtime_config(config_path)["paths"]["filejsonrst_path"], config_path=config_path, override_test_mode=override_test_mode)
    return {"status": "success" if results else "fail", "results": results}


def build_argument_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="文档信息提取系统")
    parser.add_argument("--config", default=str(PROJECT_ROOT / "config.json"))
    parser.add_argument("--payload-json", help="单组请求 JSON 字符串")
    parser.add_argument("--payload-file", help="单组请求 JSON 文件")
    parser.add_argument("--filejsonrst", help="批量任务文件路径")
    parser.add_argument("--stdin-json", action="store_true", help="从标准输入读取单组请求 JSON")
    parser.add_argument("--test-mode", choices=("true", "false"))
    return parser


def parse_cli_payload(args: argparse.Namespace) -> dict[str, Any] | None:
    if args.payload_json:
        return json.loads(args.payload_json)
    if args.payload_file:
        return json.loads(Path(args.payload_file).read_text(encoding="utf-8"))
    if args.stdin_json or (not sys.stdin.isatty() and not args.filejsonrst):
        content = sys.stdin.read().strip()
        if content:
            return json.loads(content)
    return None


def main(argv: list[str] | None = None) -> int:
    args = build_argument_parser().parse_args(argv)
    override_test_mode = None if args.test_mode is None else args.test_mode == "true"
    payload = parse_cli_payload(args)
    if payload is not None:
        result = process_request(payload, config_path=args.config, persist_result=False, override_test_mode=override_test_mode)
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0 if result.get("status") == "success" else 1
    config = load_runtime_config(args.config)
    results = process_manifest(args.filejsonrst or config["paths"]["filejsonrst_path"], config_path=args.config, override_test_mode=override_test_mode)
    print(json.dumps(results, ensure_ascii=False, indent=2))
    return 0 if results and all(item.get("status") == "success" for item in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
