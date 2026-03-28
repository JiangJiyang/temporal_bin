from __future__ import annotations

import json
import re
import subprocess
import sys
import uuid
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Sequence

import fitz
import numpy as np
import pandas as pd
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_text_splitters import MarkdownHeaderTextSplitter
from openai import OpenAI
from rapidocr_onnxruntime import RapidOCR
from settings import ConversionSettings, ModelSettings, OcrSettings, RoutingSettings

try:
    import win32com.client as win32  # type: ignore[import-not-found]
except Exception:  # pragma: no cover
    win32 = None


FONT_SIZE_MAP = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
    "六号": 7.5,
    "小六": 6.5,
    "七号": 5.5,
    "八号": 5,
}

HEADING_PATTERNS = (
    re.compile(r"^第[一二三四五六七八九十百0-9]+[章节部分篇]"),
    re.compile(r"^[一二三四五六七八九十百]+[、.]"),
    re.compile(r"^\d+(\.\d+)+"),
    re.compile(r"^\d+[、.]"),
)


@dataclass(frozen=True)
class PipelineResult:
    combined_markdown: str
    section_map: dict[str, str]
    layout_index: dict[str, Any]
    selected_branch_label: str
    route_confidence: float
    route_reason: str
    file_summaries: list[dict[str, str]]


def _is_windows() -> bool:
    return sys.platform.startswith("win")


def _normalize_rect(x0: float, y0: float, x1: float, y1: float, width: float, height: float) -> dict[str, float]:
    if width <= 0 or height <= 0:
        return {"x0": 0.0, "y0": 0.0, "x1": 1.0, "y1": 1.0}
    return {
        "x0": max(0.0, min(1.0, round(x0 / width, 6))),
        "y0": max(0.0, min(1.0, round(y0 / height, 6))),
        "x1": max(0.0, min(1.0, round(x1 / width, 6))),
        "y1": max(0.0, min(1.0, round(y1 / height, 6))),
    }


def _build_virtual_layout_lines(source_file: str, markdown_text: str) -> list[dict[str, Any]]:
    lines: list[dict[str, Any]] = []
    visible_lines = [line.strip() for line in markdown_text.splitlines() if line.strip()]
    page_size = 40
    for index, line in enumerate(visible_lines):
        page = index // page_size + 1
        offset = index % page_size
        y0 = min(0.95, 0.04 + offset * 0.022)
        y1 = min(0.98, y0 + 0.018)
        lines.append(
            {
                "source_file": source_file,
                "page": page,
                "text": line.lstrip("# ").strip(),
                "rect": {"x0": 0.06, "y0": round(y0, 6), "x1": 0.94, "y1": round(y1, 6)},
                "heading_level": 1 if line.startswith("# ") else 2 if line.startswith("## ") else 0,
            }
        )
    return lines


def _clean_json_payload(text: str) -> dict[str, Any]:
    cleaned = text.strip().strip("`").strip()
    if cleaned.lower().startswith("json"):
        cleaned = cleaned[4:].lstrip()
    match = re.search(r"\{.*\}", cleaned, re.S)
    payload = cleaned if cleaned.startswith("{") else (match.group(0) if match else cleaned)
    parsed = json.loads(payload)
    if not isinstance(parsed, dict):
        raise ValueError("Expected a JSON object.")
    return parsed


def _create_client(settings: ModelSettings) -> OpenAI:
    base_url = settings.base_url.strip()
    if base_url and not base_url.startswith(("http://", "https://")):
        base_url = f"https://{base_url}"
    return OpenAI(api_key=settings.api_key, base_url=base_url or None, timeout=settings.timeout)


def _call_model(
    settings: ModelSettings,
    *,
    system_prompt: str,
    user_prompt: str,
    temperature: float = 0,
) -> str:
    client = _create_client(settings)
    response = client.chat.completions.create(
        model=settings.name,
        temperature=temperature,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )
    return (response.choices[0].message.content or "").strip()


def _run_soffice_convert(source_path: Path, out_dir: Path, target_format: str, soffice_command: str) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    command = [
        soffice_command,
        "--headless",
        "--convert-to",
        target_format,
        "--outdir",
        str(out_dir),
        str(source_path),
    ]
    completed = subprocess.run(command, check=True, capture_output=True, text=True)
    target_suffix = "." + target_format.split(":")[0].split(";")[0]
    target_path = out_dir / f"{source_path.stem}{target_suffix}"
    if not target_path.exists():
        raise RuntimeError(f"LibreOffice conversion failed: {completed.stdout or completed.stderr}")
    return target_path.resolve()


def _convert_doc_to_docx(source_path: Path, temp_dir: Path, conversion: ConversionSettings) -> Path:
    if _is_windows() and conversion.use_office_com_on_windows and win32 is not None:
        target = temp_dir / f"{source_path.stem}_{uuid.uuid4().hex}.docx"
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        document = None
        try:
            document = word.Documents.Open(str(source_path.resolve()))
            document.SaveAs(str(target.resolve()), FileFormat=16)
            return target.resolve()
        finally:
            if document is not None:
                document.Close(False)
            word.Quit()
    return _run_soffice_convert(source_path, temp_dir, "docx", conversion.soffice_command)


def _convert_xls_to_xlsx(source_path: Path, temp_dir: Path, conversion: ConversionSettings) -> Path:
    if _is_windows() and conversion.use_office_com_on_windows and win32 is not None:
        target = temp_dir / f"{source_path.stem}_{uuid.uuid4().hex}.xlsx"
        excel = win32.Dispatch("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        workbook = None
        try:
            workbook = excel.Workbooks.Open(str(source_path.resolve()))
            workbook.SaveAs(str(target.resolve()), FileFormat=51)
            return target.resolve()
        finally:
            if workbook is not None:
                workbook.Close(False)
            excel.Quit()
    return _run_soffice_convert(source_path, temp_dir, "xlsx", conversion.soffice_command)


def _export_to_pdf(source_path: Path, temp_dir: Path, conversion: ConversionSettings) -> Path:
    if _is_windows() and conversion.use_office_com_on_windows and win32 is not None:
        suffix = source_path.suffix.lower()
        target = temp_dir / f"{source_path.stem}_{uuid.uuid4().hex}.pdf"
        if suffix in {".doc", ".docx"}:
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            document = None
            try:
                document = word.Documents.Open(str(source_path.resolve()))
                document.ExportAsFixedFormat(str(target.resolve()), 17)
                return target.resolve()
            finally:
                if document is not None:
                    document.Close(False)
                word.Quit()
        if suffix in {".xls", ".xlsx"}:
            excel = win32.Dispatch("Excel.Application")
            try:
                excel.Visible = False
            except Exception:
                pass
            try:
                excel.DisplayAlerts = False
            except Exception:
                pass
            workbook = None
            try:
                workbook = excel.Workbooks.Open(str(source_path.resolve()))
                workbook.ExportAsFixedFormat(0, str(target.resolve()))
                return target.resolve()
            finally:
                if workbook is not None:
                    workbook.Close(False)
                excel.Quit()
    return _run_soffice_convert(source_path, temp_dir, "pdf", conversion.soffice_command)


def _font_size_pt(paragraph: Paragraph) -> float | None:
    if not paragraph.runs:
        return None
    for run in paragraph.runs:
        if not run.font or not run.font.size:
            continue
        if hasattr(run.font.size, "pt"):
            return float(run.font.size.pt)
        size_str = str(run.font.size)
        for chinese_size, pt_value in FONT_SIZE_MAP.items():
            if chinese_size in size_str:
                return float(pt_value)
    return None


def _is_bold(paragraph: Paragraph) -> bool:
    return any(run.font and run.font.bold is True for run in paragraph.runs)


def _title_level(paragraph: Paragraph) -> int | None:
    text = paragraph.text.strip()
    if not text:
        return None
    if not _is_bold(paragraph):
        return 1 if any(pattern.match(text) for pattern in HEADING_PATTERNS) else None
    font_pt = _font_size_pt(paragraph)
    if font_pt is None:
        return 1 if any(pattern.match(text) for pattern in HEADING_PATTERNS) else None
    if 14.5 <= font_pt <= 15.5:
        return 1
    if 13.0 <= font_pt <= 14.5 and len(text) <= 80:
        return 2
    return 1 if any(pattern.match(text) for pattern in HEADING_PATTERNS) else None


def _paragraph_markdown(paragraph: Paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""
    level = _title_level(paragraph)
    return f"{'#' * level} {text}" if level else text


def _table_markdown(table: Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = ["\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).replace("\n", "<br>") for cell in row.cells]
        if any(cell for cell in cells):
            rows.append(cells)
    if not rows:
        return ""
    header = rows[0]
    lines = ["| " + " | ".join(header) + " |", "|" + "|".join(["---"] * len(header)) + "|"]
    for row in rows[1:]:
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(padded[: len(header)]) + " |")
    return "\n".join(lines)


def _word_to_markdown(file_path: Path, temp_dir: Path, conversion: ConversionSettings) -> str:
    actual_path = file_path
    if file_path.suffix.lower() == ".doc":
        actual_path = _convert_doc_to_docx(file_path, temp_dir, conversion)
    document = Document(BytesIO(actual_path.read_bytes()))
    parts = [part for part in (_paragraph_markdown(p) for p in document.paragraphs) if part]
    parts.extend(part for part in (_table_markdown(t) for t in document.tables) if part)
    return "\n\n".join(parts).strip()


def _excel_to_sheet_map(file_path: Path, temp_dir: Path, conversion: ConversionSettings) -> dict[str, str]:
    actual_path = file_path
    if file_path.suffix.lower() == ".xls":
        actual_path = _convert_xls_to_xlsx(file_path, temp_dir, conversion)
    data = pd.read_excel(actual_path, sheet_name=None, engine="openpyxl")
    result: dict[str, str] = {}
    for sheet_name, frame in data.items():
        frame = frame.fillna("")
        headers = [str(item) for item in frame.columns.tolist()]
        lines = [
            f"## {sheet_name}",
            "",
            f"Rows: {len(frame)}, Columns: {len(headers)}",
            "",
            "| " + " | ".join(headers) + " |",
            "|" + "|".join(["---"] * len(headers)) + "|",
        ]
        for _, row in frame.iterrows():
            values = [str(row[h]).replace("|", "\\|").replace("\n", "<br>") for h in frame.columns]
            lines.append("| " + " | ".join(values) + " |")
        result[str(sheet_name)] = "\n".join(lines)
    return result


def _normalize_section_key(value: str) -> str:
    return re.sub(r"\d", "", value).replace(" ", "").replace(".", "").strip()


def split_sections(markdown_text: str) -> dict[str, str]:
    total = ""
    for line in markdown_text.splitlines():
        if line.lstrip().startswith("#"):
            start_count = 0
            for char in line:
                if char == "#":
                    start_count += 1
                else:
                    break
            total += "#" * max(1, line.count(".") - start_count + 1) + line + "\n"
        else:
            total += line + "\n"
    splitter = MarkdownHeaderTextSplitter(headers_to_split_on=[("#", "Header 1"), ("##", "Header 2")])
    sections: dict[str, str] = {}
    for doc in splitter.split_text(total):
        for value in doc.metadata.values():
            key = _normalize_section_key(str(value))
            if not key:
                continue
            sections[key] = sections.get(key, key) + ("\n" if key in sections else "") + doc.page_content
    sections["封面"] = total[:80]
    return sections


def merge_text_maps(*maps: dict[str, str]) -> dict[str, str]:
    merged: dict[str, str] = {}
    for current in maps:
        for key, value in current.items():
            if not value:
                continue
            merged[key] = f"{merged[key]}\n\n{value}" if key in merged and merged[key] != value else value
    return merged


def _heuristic_heading_level(text: str, font_size: float | None = None, is_bold: bool = False) -> int:
    stripped = text.strip()
    if not stripped:
        return 0
    if any(pattern.match(stripped) for pattern in HEADING_PATTERNS):
        return 2 if len(stripped) <= 80 else 0
    if is_bold and font_size and font_size >= 14 and len(stripped) <= 80:
        return 1
    if font_size and font_size >= 15 and len(stripped) <= 80:
        return 1
    if font_size and font_size >= 13 and len(stripped) <= 80:
        return 2
    return 0


def _get_ocr_engine() -> RapidOCR:
    return RapidOCR()


def _pdf_page_to_lines(page: fitz.Page, source_file: str) -> list[dict[str, Any]]:
    page_dict = page.get_text("dict")
    width = float(page.rect.width)
    height = float(page.rect.height)
    lines: list[dict[str, Any]] = []
    for block in page_dict.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            text = "".join(str(span.get("text", "")) for span in spans).strip()
            if not text:
                continue
            font_size = max((float(span.get("size", 0)) for span in spans), default=0.0)
            is_bold = any(("Bold" in str(span.get("font", ""))) or (int(span.get("flags", 0)) & 16) for span in spans)
            x0, y0, x1, y1 = map(float, line.get("bbox", (0, 0, width, height)))
            lines.append(
                {
                    "source_file": source_file,
                    "page": page.number + 1,
                    "text": text,
                    "rect": _normalize_rect(x0, y0, x1, y1, width, height),
                    "heading_level": _heuristic_heading_level(text, font_size=font_size, is_bold=is_bold),
                }
            )
    return lines


def _ocr_page_to_lines(page: fitz.Page, source_file: str, ocr: OcrSettings) -> list[dict[str, Any]]:
    zoom = max(ocr.dpi / 72.0, 1.0)
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
    engine = _get_ocr_engine()
    result, _ = engine(img_array)
    if not result:
        return []
    lines: list[dict[str, Any]] = []
    for item in result:
        points, text, _score = item
        if not str(text).strip():
            continue
        xs = [float(point[0]) for point in points]
        ys = [float(point[1]) for point in points]
        lines.append(
            {
                "source_file": source_file,
                "page": page.number + 1,
                "text": str(text).strip(),
                "rect": _normalize_rect(min(xs), min(ys), max(xs), max(ys), float(pix.width), float(pix.height)),
                "heading_level": _heuristic_heading_level(str(text)),
            }
        )
    lines.sort(key=lambda item: (item["rect"]["y0"], item["rect"]["x0"]))
    return lines


def _extract_pdf_lines(pdf_path: Path, ocr: OcrSettings) -> list[dict[str, Any]]:
    lines: list[dict[str, Any]] = []
    document = fitz.open(pdf_path)
    try:
        for page in document:
            if page.get_text("text").strip():
                lines.extend(_pdf_page_to_lines(page, pdf_path.name))
            elif ocr.enabled:
                lines.extend(_ocr_page_to_lines(page, pdf_path.name, ocr))
    finally:
        document.close()
    return lines


def _lines_to_markdown(file_name: str, lines: list[dict[str, Any]]) -> str:
    chunks: list[str] = [f"# 文件：{file_name}"]
    current_page = None
    for line in lines:
        if line["page"] != current_page:
            current_page = line["page"]
            chunks.extend(["", f"## 第{current_page}页", ""])
        prefix = "#" * line["heading_level"] + " " if line["heading_level"] else ""
        chunks.append(f"{prefix}{line['text']}")
    return "\n".join(chunks).strip()


def _build_pdf_markdown_and_layout(pdf_path: Path, ocr: OcrSettings) -> tuple[str, list[dict[str, Any]]]:
    lines = _extract_pdf_lines(pdf_path, ocr)
    return _lines_to_markdown(pdf_path.name, lines), lines


def _choose_routing_model(
    routing: RoutingSettings,
    qwen: ModelSettings,
    deepseek: ModelSettings,
) -> ModelSettings:
    base = qwen if routing.model_key == "qwen" else deepseek
    if routing.model_name:
        return ModelSettings(
            provider=base.provider,
            name=routing.model_name,
            api_key=base.api_key,
            base_url=base.base_url,
            timeout=base.timeout,
        )
    return base


def _fallback_branch_label(question: str, file_summaries: list[dict[str, str]], branches: Sequence[Any]) -> str:
    haystack = " ".join([question] + [item.get("file_name", "") + " " + item.get("preview", "") for item in file_summaries]).lower()
    best_label = getattr(branches[0], "label")
    best_score = -1
    for branch in branches:
        score = 0
        for keyword in getattr(branch, "match_keywords", []):
            if keyword.lower() in haystack:
                score += max(1, len(keyword))
        if score > best_score:
            best_score = score
            best_label = getattr(branch, "label")
    return best_label


def _route_branch(
    *,
    question: str,
    file_summaries: list[dict[str, str]],
    branches: Sequence[Any],
    routing: RoutingSettings,
    routing_model: ModelSettings,
) -> tuple[str, float, str]:
    branch_options = [
        {
            "label": getattr(branch, "label"),
            "match_keywords": list(getattr(branch, "match_keywords", [])),
            "model_key": getattr(branch, "model_key"),
            "field_samples": list(getattr(branch, "field_map", {}).keys())[:12],
        }
        for branch in branches
    ]
    user_prompt = json.dumps(
        {
            "question": question,
            "uploaded_files": file_summaries,
            "candidate_branches": branch_options,
        },
        ensure_ascii=False,
        indent=2,
    )
    raw_output = _call_model(
        routing_model,
        system_prompt=routing.system_prompt,
        user_prompt=user_prompt,
        temperature=routing.temperature,
    )
    payload = _clean_json_payload(raw_output)
    label = str(payload.get("branch_label", "")).strip()
    confidence = float(payload.get("confidence", 0))
    reason = str(payload.get("reason", "")).strip()
    valid_labels = {getattr(branch, "label") for branch in branches}
    if label not in valid_labels:
        raise ValueError(f"Invalid branch label returned by router: {label}")
    return label, confidence, reason


def ingest_and_route(
    *,
    files: Sequence[Path],
    question: str,
    temp_dir: Path,
    qwen: ModelSettings,
    deepseek: ModelSettings,
    routing: RoutingSettings,
    ocr: OcrSettings,
    conversion: ConversionSettings,
    branches: Sequence[Any],
) -> PipelineResult:
    doc_parts: list[str] = []
    section_maps: list[dict[str, str]] = []
    layout_lines: list[dict[str, Any]] = []
    file_summaries: list[dict[str, str]] = []

    for file_path in files:
        suffix = file_path.suffix.lower()
        preview_markdown = ""
        if suffix in {".doc", ".docx"}:
            preview_markdown = _word_to_markdown(file_path, temp_dir, conversion)
            doc_parts.append(f"# 文件：{file_path.name}\n\n{preview_markdown}")
            section_maps.append(split_sections(preview_markdown))
            layout_pdf = _export_to_pdf(file_path, temp_dir, conversion)
            _layout_markdown, lines = _build_pdf_markdown_and_layout(layout_pdf, ocr)
            layout_lines.extend(lines)
        elif suffix in {".xls", ".xlsx"}:
            sheet_map = _excel_to_sheet_map(file_path, temp_dir, conversion)
            preview_markdown = "\n\n".join(sheet_map.values()).strip()
            doc_parts.append(f"# 文件：{file_path.name}\n\n{preview_markdown}")
            section_maps.append(sheet_map)
            try:
                layout_pdf = _export_to_pdf(file_path, temp_dir, conversion)
                _layout_markdown, lines = _build_pdf_markdown_and_layout(layout_pdf, ocr)
                layout_lines.extend(lines or _build_virtual_layout_lines(file_path.name, preview_markdown))
            except Exception:
                layout_lines.extend(_build_virtual_layout_lines(file_path.name, preview_markdown))
        elif suffix == ".pdf":
            preview_markdown, lines = _build_pdf_markdown_and_layout(file_path, ocr)
            doc_parts.append(preview_markdown)
            section_maps.append(split_sections(preview_markdown))
            layout_lines.extend(lines)
        else:
            raise ValueError(f"Unsupported file type: {file_path.name}")

        file_summaries.append(
            {
                "file_name": file_path.name,
                "file_type": suffix.lstrip("."),
                "preview": preview_markdown[: routing.max_chars],
            }
        )

    combined_markdown = "\n\n".join(part for part in doc_parts if part).strip()
    section_map = merge_text_maps(*section_maps)

    selected_branch_label = ""
    route_confidence = 0.0
    route_reason = ""
    try:
        routing_model = _choose_routing_model(routing, qwen, deepseek)
        selected_branch_label, route_confidence, route_reason = _route_branch(
            question=question,
            file_summaries=file_summaries,
            branches=branches,
            routing=routing,
            routing_model=routing_model,
        )
    except Exception as exc:
        if not routing.fallback_to_keyword:
            raise
        selected_branch_label = _fallback_branch_label(question, file_summaries, branches)
        route_confidence = 0.0
        route_reason = f"fallback_to_keyword: {exc}"

    return PipelineResult(
        combined_markdown=combined_markdown,
        section_map=section_map,
        layout_index={"lines": layout_lines},
        selected_branch_label=selected_branch_label,
        route_confidence=route_confidence,
        route_reason=route_reason,
        file_summaries=file_summaries,
    )
